# -*- coding: utf-8 -*-
"""Generate docx syllabus for Data Structures and Algorithms course."""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from generate_syllabus_docx import SyllabusDocx
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = SyllabusDocx()

# Title
doc.add_title('\u300a数据结构与算法\u300b课程教学大纲')

# ===== Section 1 =====
doc.add_h1('一、课程基本信息')
doc.add_table(
    ['项目', '内容'],
    [
        [('课程编码', WD_ALIGN_PARAGRAPH.CENTER, False), ('ZB2304012', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('课程名称', WD_ALIGN_PARAGRAPH.CENTER, False), ('数据结构与算法', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('课程英文名称', WD_ALIGN_PARAGRAPH.CENTER, False), ('Data Structures and Algorithms', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('适用专业', WD_ALIGN_PARAGRAPH.CENTER, False), ('智能科学与技术', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('课程性质', WD_ALIGN_PARAGRAPH.CENTER, False), ('学科基础必修课', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('开课学期', WD_ALIGN_PARAGRAPH.CENTER, False), ('第3学期', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('学分', WD_ALIGN_PARAGRAPH.CENTER, False), ('4.0', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('总学时', WD_ALIGN_PARAGRAPH.CENTER, False), ('64', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('讲课学时', WD_ALIGN_PARAGRAPH.CENTER, False), ('48', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('实验/上机学时', WD_ALIGN_PARAGRAPH.CENTER, False), ('16', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('考核方式', WD_ALIGN_PARAGRAPH.CENTER, False), ('考试\uff08闭卷笔试+实验考核\uff09', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('先修课程', WD_ALIGN_PARAGRAPH.CENTER, False), ('高级语言程序设计\uff08C/C++\uff09、离散数学', WD_ALIGN_PARAGRAPH.LEFT, False)],
    ],
    [3.0, 11.6]
)

# ===== Section 2 =====
doc.add_h1('二、课程性质与任务')
doc.add_h2('\uff08一\uff09课程性质')
doc.add_body('本课程是智能科学与技术专业的一门学科基础必修课，是计算机科学与技术领域的核心基础课程。课程系统讲授线性表、栈与队列、树与二叉树、图、查找与排序等经典数据结构的基本概念、逻辑结构、存储结构及基本操作算法，培养学生根据实际问题选择合适数据结构与算法设计范式的能力。')
doc.add_h2('\uff08二\uff09课程任务')
doc.add_body('本课程的任务是使学生掌握各种经典数据结构的基本概念、原理和实现方法，具备运用抽象数据类型\uff08ADT\uff09进行问题建模的能力，能够分析算法的时间复杂度与空间复杂度，并能综合运用数据结构与算法知识解决智能科学与技术领域中的实际问题，为后续专业课程的学习及从事智能系统研发奠定坚实的算法基础。')

# ===== Section 3 =====
doc.add_h1('三、课程目标与毕业要求支撑关系')
doc.add_h2('\uff08一\uff09课程目标')
doc.add_table(
    ['编号', '课程目标', '对应毕业要求'],
    [
        [('目标1', WD_ALIGN_PARAGRAPH.CENTER, False), ('掌握线性表、栈、队列、树、图、查找表、排序等基本数据结构的逻辑结构、存储结构及基本操作算法', WD_ALIGN_PARAGRAPH.LEFT, False), ('毕业要求1\uff08工程知识\uff09', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('目标2', WD_ALIGN_PARAGRAPH.CENTER, False), ('能够针对实际问题进行数据结构选型与算法设计，具备抽象数据类型建模与算法复杂度分析能力', WD_ALIGN_PARAGRAPH.LEFT, False), ('毕业要求2\uff08问题分析\uff09', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('目标3', WD_ALIGN_PARAGRAPH.CENTER, False), ('能够综合运用数据结构与算法知识，设计和实现满足特定需求的计算解决方案', WD_ALIGN_PARAGRAPH.LEFT, False), ('毕业要求3\uff08设计/开发解决方案\uff09', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('目标4', WD_ALIGN_PARAGRAPH.CENTER, False), ('熟练使用C/C++或Python语言实现基本数据结构与算法，能利用现代开发调试工具验证算法正确性', WD_ALIGN_PARAGRAPH.LEFT, False), ('毕业要求5\uff08使用现代工具\uff09', WD_ALIGN_PARAGRAPH.CENTER, False)],
    ],
    [1.5, 8.0, 5.1]
)

doc.add_h2('\uff08二\uff09毕业要求支撑矩阵')
doc.add_table(
    ['毕业要求', '目标1', '目标2', '目标3', '目标4'],
    [
        [('要求1\uff1a工程知识', WD_ALIGN_PARAGRAPH.CENTER, False), ('H', WD_ALIGN_PARAGRAPH.CENTER, False), ('M', WD_ALIGN_PARAGRAPH.CENTER, False), ('', WD_ALIGN_PARAGRAPH.CENTER, False), ('', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('要求2\uff1a问题分析', WD_ALIGN_PARAGRAPH.CENTER, False), ('M', WD_ALIGN_PARAGRAPH.CENTER, False), ('H', WD_ALIGN_PARAGRAPH.CENTER, False), ('M', WD_ALIGN_PARAGRAPH.CENTER, False), ('', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('要求3\uff1a设计/开发解决方案', WD_ALIGN_PARAGRAPH.CENTER, False), ('', WD_ALIGN_PARAGRAPH.CENTER, False), ('M', WD_ALIGN_PARAGRAPH.CENTER, False), ('H', WD_ALIGN_PARAGRAPH.CENTER, False), ('M', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('要求5\uff1a使用现代工具', WD_ALIGN_PARAGRAPH.CENTER, False), ('', WD_ALIGN_PARAGRAPH.CENTER, False), ('', WD_ALIGN_PARAGRAPH.CENTER, False), ('L', WD_ALIGN_PARAGRAPH.CENTER, False), ('H', WD_ALIGN_PARAGRAPH.CENTER, False)],
    ],
    [4.0, 2.0, 2.0, 2.0, 2.0]
)
doc.add_body('\uff08H=高支撑，M=中支撑，L=弱支撑\uff09')

# ===== Section 4 =====
doc.add_h1('四、教学内容与学时分配')
doc.add_body('本课程共64学时，其中讲课48学时，实验/上机16学时。教学内容按如下章节组织：')
doc.add_h2('第一章 绪论\uff084学时，讲课4学时\uff09')
doc.add_body('教学内容：数据结构的基本概念\uff08数据、数据元素、数据对象、数据结构\uff09；逻辑结构与存储结构；抽象数据类型\uff08ADT\uff09的概念与表示；算法与算法分析\uff08时间复杂度和空间复杂度\uff09；智能科学与技术领域中的数据结构应用概览。')
doc.add_body_mixed([('重点：', True), ('数据结构的三要素\uff08逻辑结构、存储结构、运算\uff09，算法时间复杂度分析', False)])
doc.add_body_mixed([('难点：', True), ('渐进时间复杂度的推导与比较', False)])
doc.add_body('思政融入点：从\u201c程序=数据结构+算法\u201d的经典命题引出计算思维在认识世界中的普遍性，培养学生的科学思维方法。')

doc.add_h2('第二章 线性表\uff088学时，讲课6学时+上机2学时\uff09')
doc.add_body('教学内容：线性表的定义与基本操作；顺序表的实现\uff08静态分配与动态分配\uff09；单链表、双链表、循环链表的实现；顺序表与链表的比较与选用；线性表的应用：一元多项式表示与运算。')
doc.add_body_mixed([('重点：', True), ('顺序表和链表的基本操作实现\uff08插入、删除、查找\uff09', False)])
doc.add_body_mixed([('难点：', True), ('链表操作的边界条件处理与指针操作', False)])
doc.add_body('实验1：线性表的基本操作实现与性能对比\uff082学时\uff09')

doc.add_h2('第三章 栈与队列\uff088学时，讲课6学时+上机2学时\uff09')
doc.add_body('教学内容：栈的定义、顺序栈与链栈的实现；栈的应用：表达式求值、括号匹配、递归转非递归；队列的定义、顺序队列与链队列的实现；循环队列的设计；队列的应用：层次遍历、缓冲区管理。')
doc.add_body_mixed([('重点：', True), ('栈和队列的ADT定义与基本操作；循环队列的队空队满判定', False)])
doc.add_body_mixed([('难点：', True), ('栈实现表达式求值的算符优先算法，递归的工作栈原理', False)])
doc.add_body('实验2：基于栈的表达式求值或迷宫求解\uff082学时\uff09')

doc.add_h2('第四章 串、数组与广义表\uff084学时，讲课4学时\uff09')
doc.add_body('教学内容：串的定义与基本操作\uff08模式匹配\uff09；KMP算法；数组的存储结构；矩阵的压缩存储\uff08对称矩阵、三角矩阵、稀疏矩阵\uff09；广义表的概念。')
doc.add_body_mixed([('重点：', True), ('KMP算法的next数组求解，特殊矩阵的压缩存储', False)])
doc.add_body_mixed([('难点：', True), ('KMP算法的思想与nextval数组优化', False)])

doc.add_h2('第五章 树与二叉树\uff0810学时，讲课8学时+上机2学时\uff09')
doc.add_body('教学内容：树的基本概念与术语；二叉树的定义、性质与存储结构；二叉树的遍历\uff08前序、中序、后序、层次\uff09；线索二叉树；树与森林的转换与遍历；哈夫曼树与哈夫曼编码；并查集。')
doc.add_body_mixed([('重点：', True), ('二叉树的递归遍历算法，哈夫曼树的构造与编码', False)])
doc.add_body_mixed([('难点：', True), ('非递归遍历算法，线索二叉树的构造', False)])
doc.add_body('实验3：二叉树遍历与哈夫曼编码实现\uff082学时\uff09')

doc.add_h2('第六章 图\uff0810学时，讲课8学时+上机2学时\uff09')
doc.add_body('教学内容：图的定义与术语\uff08有向图、无向图、连通性\uff09；图的存储结构\uff08邻接矩阵、邻接表\uff09；图的遍历\uff08深度优先搜索DFS、广度优先搜索BFS\uff09；最小生成树\uff08Prim算法、Kruskal算法\uff09；最短路径\uff08Dijkstra算法、Floyd算法\uff09；拓扑排序与关键路径。')
doc.add_body_mixed([('重点：', True), ('DFS与BFS的算法实现，Dijkstra最短路径算法', False)])
doc.add_body_mixed([('难点：', True), ('关键路径算法的理解与实现', False)])
doc.add_body('实验4：图的遍历与最短路径算法实现\uff082学时\uff09')

doc.add_h2('第七章 查找\uff088学时，讲课6学时+上机2学时\uff09')
doc.add_body('教学内容：查找的基本概念与性能指标\uff08ASL\uff09；顺序查找与折半查找；二叉排序树\uff08BST\uff09与平衡二叉树\uff08AVL\uff09；哈希表与哈希冲突处理；B树与B+树的基本概念。')
doc.add_body_mixed([('重点：', True), ('折半查找算法，二叉排序树的插入与删除，哈希表的构造与查找', False)])
doc.add_body_mixed([('难点：', True), ('AVL树的平衡调整，哈希冲突的链地址法与开放定址法', False)])
doc.add_body('实验5：哈希表设计与性能分析\uff082学时\uff09')

doc.add_h2('第八章 排序\uff088学时，讲课6学时+上机2学时\uff09')
doc.add_body('教学内容：排序的基本概念与分类\uff08内排序/外排序、稳定/不稳定\uff09；插入排序\uff08直接插入、希尔排序\uff09；交换排序\uff08冒泡排序、快速排序\uff09；选择排序\uff08简单选择、堆排序\uff09；归并排序\uff08二路归并\uff09；基数排序；各种排序算法的比较与选择。')
doc.add_body_mixed([('重点：', True), ('快速排序、堆排序、归并排序的算法思想与实现', False)])
doc.add_body_mixed([('难点：', True), ('堆的调整算法与堆排序过程', False)])
doc.add_body('实验6：多种排序算法性能对比实验\uff082学时\uff09')

doc.add_h2('实验环节汇总\uff0816学时\uff09')
doc.add_table(
    ['序号', '实验项目', '学时', '实验类型', '支撑目标'],
    [
        [('1', WD_ALIGN_PARAGRAPH.CENTER, False), ('线性表基本操作实现与性能对比', WD_ALIGN_PARAGRAPH.LEFT, False), ('2', WD_ALIGN_PARAGRAPH.CENTER, False), ('验证性', WD_ALIGN_PARAGRAPH.CENTER, False), ('目标1、目标4', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('2', WD_ALIGN_PARAGRAPH.CENTER, False), ('基于栈的表达式求值或迷宫求解', WD_ALIGN_PARAGRAPH.LEFT, False), ('2', WD_ALIGN_PARAGRAPH.CENTER, False), ('设计性', WD_ALIGN_PARAGRAPH.CENTER, False), ('目标2、目标4', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('3', WD_ALIGN_PARAGRAPH.CENTER, False), ('二叉树遍历与哈夫曼编码实现', WD_ALIGN_PARAGRAPH.LEFT, False), ('2', WD_ALIGN_PARAGRAPH.CENTER, False), ('设计性', WD_ALIGN_PARAGRAPH.CENTER, False), ('目标1、目标3', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('4', WD_ALIGN_PARAGRAPH.CENTER, False), ('图的遍历与最短路径算法实现', WD_ALIGN_PARAGRAPH.LEFT, False), ('2', WD_ALIGN_PARAGRAPH.CENTER, False), ('设计性', WD_ALIGN_PARAGRAPH.CENTER, False), ('目标2、目标3', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('5', WD_ALIGN_PARAGRAPH.CENTER, False), ('哈希表设计与性能分析', WD_ALIGN_PARAGRAPH.LEFT, False), ('2', WD_ALIGN_PARAGRAPH.CENTER, False), ('综合性', WD_ALIGN_PARAGRAPH.CENTER, False), ('目标2、目标4', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('6', WD_ALIGN_PARAGRAPH.CENTER, False), ('多种排序算法性能对比实验', WD_ALIGN_PARAGRAPH.LEFT, False), ('2', WD_ALIGN_PARAGRAPH.CENTER, False), ('综合性', WD_ALIGN_PARAGRAPH.CENTER, False), ('目标1、目标4', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('7', WD_ALIGN_PARAGRAPH.CENTER, False), ('综合项目：学生成绩管理系统或通讯录管理系统', WD_ALIGN_PARAGRAPH.LEFT, False), ('4', WD_ALIGN_PARAGRAPH.CENTER, False), ('综合性', WD_ALIGN_PARAGRAPH.CENTER, False), ('目标2、目标3、目标4', WD_ALIGN_PARAGRAPH.CENTER, False)],
    ],
    [1.0, 5.5, 1.2, 1.8, 2.5]
)

# ===== Section 5 =====
doc.add_h1('五、教学方法与手段')
doc.add_h2('\uff08一\uff09教学方法')
doc.add_body('本课程采用\u201c理论讲授+实验实践+创新拓展\u201d三维融合教学模式：')
doc.add_body('1. 理论讲授\uff0848学时\uff09：以课堂讲授为主，结合板书与多媒体演示，注重算法逻辑的推演与可视化展示。每章采用\u201c问题驱动\u201d教学法\u2014\u2014从实际问题出发引出数据结构的必要性，讲解原理后回归应用。')
doc.add_body('2. 实验实践\uff0816学时\uff09：通过6个单项实验和1个综合项目，强化编程实现能力。实验采用\u201c任务驱动+代码走查\u201d方式逐步递进。')
doc.add_body('3. 创新拓展\uff08课外\uff09：引入LeetCode/Codeforces等在线评测平台上的经典题目作为进阶练习，鼓励学生参与算法竞赛，培养创新思维。')
doc.add_h2('\uff08二\uff09教学手段')
doc.add_body('1. 多媒体教学：使用动效演示栈、队列、树、图等数据结构的动态操作过程，降低认知门槛。')
doc.add_body('2. 智慧教学平台：依托超星学习通/雨课堂等平台进行课堂互动、随堂测验和课后讨论，实现即时学情反馈。')
doc.add_body('3. 在线评测系统：利用OJ\uff08Online Judge\uff09系统布置编程实验，支持自动判题与代码质量分析。')
doc.add_body('4. AI辅助学习：引导学生使用ChatGPT/GitHub Copilot等AI工具辅助代码调试与算法理解，同时培养批判性审视AI生成代码的能力。教师利用AI分析学生作业中的常见错误模式，数据驱动改进教学。')

# ===== Section 6 =====
doc.add_h1('六、课程思政设计')
doc.add_table(
    ['思政主题', '融入章节', '融入方式', '育人成效'],
    [
        [('科学思维与计算思维', WD_ALIGN_PARAGRAPH.CENTER, False), ('第一章 绪论', WD_ALIGN_PARAGRAPH.CENTER, False), ('从\u201c程序=数据结构+算法\u201d引出计算思维的方法论意义', WD_ALIGN_PARAGRAPH.LEFT, False), ('培养学生用计算思维分析问题的习惯', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('辩证思维与工程权衡', WD_ALIGN_PARAGRAPH.CENTER, False), ('第二章 线性表', WD_ALIGN_PARAGRAPH.CENTER, False), ('顺序表vs链表的时间/空间权衡分析', WD_ALIGN_PARAGRAPH.LEFT, False), ('培养学生在约束条件下寻求最优解的工程思维', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('效率优化与节约意识', WD_ALIGN_PARAGRAPH.CENTER, False), ('第五章 树与二叉树', WD_ALIGN_PARAGRAPH.CENTER, False), ('哈夫曼编码的最小化加权路径长度思想', WD_ALIGN_PARAGRAPH.LEFT, False), ('培养学生资源节约与优化意识', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('决策智慧与工程使命', WD_ALIGN_PARAGRAPH.CENTER, False), ('第六章 图', WD_ALIGN_PARAGRAPH.CENTER, False), ('最短路径类比人生选择，结合智慧交通案例', WD_ALIGN_PARAGRAPH.LEFT, False), ('激发科技报国、服务社会的责任感', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('秩序公平与制度设计', WD_ALIGN_PARAGRAPH.CENTER, False), ('第八章 排序', WD_ALIGN_PARAGRAPH.CENTER, False), ('排序稳定性与社会治理的类比', WD_ALIGN_PARAGRAPH.LEFT, False), ('培养学生公平公正的价值观', WD_ALIGN_PARAGRAPH.LEFT, False)],
    ],
    [2.5, 2.5, 5.5, 4.1]
)

# ===== Section 7 =====
doc.add_h1('七、教材与学习资源')
doc.add_h2('\uff08一\uff09教材')
doc.add_body('1. 主教材：严蔚敏，吴伟民.《数据结构\uff08C语言版\uff09》. 北京：清华大学出版社，2021.\uff08\u201c十二五\u201d普通高等教育本科国家级规划教材\uff09')
doc.add_body('2. 实验教材：李春葆.《数据结构习题与解析\uff08C语言版\uff09》. 北京：清华大学出版社，2022.')
doc.add_h2('\uff08二\uff09参考书目')
doc.add_body('1. Thomas H. Cormen 等.《算法导论\uff08原书第4版\uff09》. 北京：机械工业出版社，2023.')
doc.add_body('2. Mark Allen Weiss.《数据结构与算法分析\u2014\u2014C语言描述\uff08原书第2版\uff09》. 北京：机械工业出版社，2020.')
doc.add_body('3. 邓俊辉.《数据结构\uff08C++语言版\xb7第4版\uff09》. 北京：清华大学出版社，2021.')
doc.add_h2('\uff08三\uff09在线资源')
doc.add_body('1. 中国大学MOOC\u2014\u2014数据结构\uff08浙江大学\uff09')
doc.add_body('2. LeetCode 数据结构专项')
doc.add_body('3. VisuAlgo 算法可视化平台')
doc.add_body('4. GeeksforGeeks Data Structures')
doc.add_body('5. 学校在线评测系统\uff08OJ\uff09')

# ===== Section 8 =====
doc.add_h1('八、考核方式与成绩评定')
doc.add_h2('\uff08一\uff09考核方式')
doc.add_table(
    ['考核环节', '占比', '考核内容', '支撑目标'],
    [
        [('平时表现', WD_ALIGN_PARAGRAPH.CENTER, False), ('10%', WD_ALIGN_PARAGRAPH.CENTER, False), ('课堂互动、随堂测验、作业完成情况', WD_ALIGN_PARAGRAPH.LEFT, False), ('目标1、目标2', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('实验报告与编程作业', WD_ALIGN_PARAGRAPH.CENTER, False), ('20%', WD_ALIGN_PARAGRAPH.CENTER, False), ('6次实验报告+综合项目代码与文档', WD_ALIGN_PARAGRAPH.LEFT, False), ('目标3、目标4', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('期中考试', WD_ALIGN_PARAGRAPH.CENTER, False), ('10%', WD_ALIGN_PARAGRAPH.CENTER, False), ('前半部分内容\uff08线性表、栈、队列、串\uff09', WD_ALIGN_PARAGRAPH.LEFT, False), ('目标1、目标2', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('期末考试', WD_ALIGN_PARAGRAPH.CENTER, False), ('50%', WD_ALIGN_PARAGRAPH.CENTER, False), ('闭卷笔试，涵盖全部内容', WD_ALIGN_PARAGRAPH.LEFT, False), ('目标1、目标2、目标3', WD_ALIGN_PARAGRAPH.CENTER, False)],
        [('课程思政表现', WD_ALIGN_PARAGRAPH.CENTER, False), ('10%', WD_ALIGN_PARAGRAPH.CENTER, False), ('思政案例分析报告或课堂讨论表现', WD_ALIGN_PARAGRAPH.LEFT, False), ('目标1', WD_ALIGN_PARAGRAPH.CENTER, False)],
    ],
    [2.5, 1.5, 5.5, 2.5]
)
doc.add_h2('\uff08二\uff09成绩评定')
doc.add_body('总成绩 = 平时表现\xd710% + 实验与编程作业\xd720% + 期中考试\xd710% + 期末考试\xd750% + 课程思政\xd710%')
doc.add_body('实验环节必须全部提交且验收通过，否则取消期末考试资格。期末考试成绩低于45分者，总成绩直接评定为不及格。')

# ===== Section 9 =====
doc.add_h1('九、教学改革说明')
doc.add_h2('\uff08一\uff09立德树人')
doc.add_body('本课程围绕\u201c计算思维塑造+工程伦理培养+科技报国情怀\u201d三条主线开展课程思政。充分利用数据结构学科发展史中的中国贡献\uff08如中文信息处理的词典数据结构创新、国产数据库的索引技术创新等\uff09激发学生的民族自信心和科技报国责任感。每章至少设置一个思政融入点，实现知识传授与价值引领的有机统一。')
doc.add_h2('\uff08二\uff09教学思想')
doc.add_body('本课程坚持\u201c学生中心、产出导向\u201d的教学理念，以OBE\uff08Outcome-Based Education\uff09为指导，反向设计教学内容与考核方式。面对AI时代挑战，课程注重培养学生在AI辅助环境下的算法设计与代码审校能力\u2014\u2014既要善用AI工具提升开发效率，更要具备独立理解、分析、优化算法的核心素养。')
doc.add_h2('\uff08三\uff09教学方法')
doc.add_body('采用\u201c理论+实践+创新\u201d三维融合教学模式：理论讲授环节注重算法逻辑的可视化推演与推导式教学；实验环节采用\u201c验证性实验\u2192设计性实验\u2192综合性项目\u201d三阶递进体系；创新拓展环节引入算法竞赛训练与科研小课题，为学有余力的学生提供提升通道。')
doc.add_h2('\uff08四\uff09教学手段')
doc.add_body('充分运用超星学习通/雨课堂等智慧教学平台实现课堂互动与学情数据采集；引入OJ自动评测系统实现编程作业的自动化批改与即时反馈；利用代码可视化工具和AI辅助编程工具降低学生入门门槛；建立学生代码错误数据库，数据驱动地优化教学重点与难点突破策略。')

# Appendix
doc.add_page_break()
doc.add_h1('编制说明与合规自查报告')
doc.add_h2('一、专业与课程匹配说明')
doc.add_table(
    ['项目', '内容'],
    [
        [('专业名称', WD_ALIGN_PARAGRAPH.CENTER, False), ('智能科学与技术', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('课程名称', WD_ALIGN_PARAGRAPH.CENTER, False), ('数据结构与算法', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('课程编码', WD_ALIGN_PARAGRAPH.CENTER, False), ('ZB2304012', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('学分/学时', WD_ALIGN_PARAGRAPH.CENTER, False), ('4.0学分/64学时', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('编制依据', WD_ALIGN_PARAGRAPH.CENTER, False), ('智能科学与技术专业2026版培养方案\uff08通用模板参考\uff09', WD_ALIGN_PARAGRAPH.LEFT, False)],
    ],
    [3.0, 11.6]
)
doc.add_h2('二、课程性质判定依据')
doc.add_body('课程以知识传授与原理讲授为主\uff0848学时理论+16学时实验\uff09；实验环节服务于理论验证与编程能力培养，非独立设计项目；课程名称不含\u201c课程设计\u201d\u201c设计\u201d等字眼。')
doc.add_body_mixed([('判定结果：理论课程', True), ('，匹配模板《1. 理论课程教学大纲的格式.docx》', False)])
doc.add_h2('三、工程认证支撑关系')
doc.add_body('基于智能科学与技术专业培养定位，本课程重点支撑以下毕业要求：毕业要求1\uff08工程知识\uff09\u2014\u2014H支撑；毕业要求2\uff08问题分析\uff09\u2014\u2014H支撑；毕业要求3\uff08设计/开发解决方案\uff09\u2014\u2014M支撑；毕业要求5\uff08使用现代工具\uff09\u2014\u2014M支撑。')
doc.add_h2('四、八大维度合规自查结果')
doc.add_table(
    ['维度', '要求', '自查结果'],
    [
        [('1. 元数据锁定', WD_ALIGN_PARAGRAPH.CENTER, False), ('编码、名称、总学时严格对齐培养方案', WD_ALIGN_PARAGRAPH.LEFT, False), ('\u2705 编码参考同类专业设置', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('2. 先修课程', WD_ALIGN_PARAGRAPH.CENTER, False), ('基于课程体系确认合理前置关系', WD_ALIGN_PARAGRAPH.LEFT, False), ('\u2705 高级语言程序设计、离散数学', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('3. 性质与任务', WD_ALIGN_PARAGRAPH.CENTER, False), ('支撑专业毕业要求指标点', WD_ALIGN_PARAGRAPH.LEFT, False), ('\u2705 明确支撑3项毕业要求', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('4. 内容与重难点', WD_ALIGN_PARAGRAPH.CENTER, False), ('高阶性、创新性、挑战度', WD_ALIGN_PARAGRAPH.LEFT, False), ('\u2705 64学时含实验16学时', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('5. 能力培养', WD_ALIGN_PARAGRAPH.CENTER, False), ('创新思维+实践能力+数字素养', WD_ALIGN_PARAGRAPH.LEFT, False), ('\u2705 4项课程目标对应具体载体', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('6. 教材与资源', WD_ALIGN_PARAGRAPH.CENTER, False), ('优先国规教材，可访问线上资源', WD_ALIGN_PARAGRAPH.LEFT, False), ('\u2705 严蔚敏版+多项资源', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('7. 考核与成绩', WD_ALIGN_PARAGRAPH.CENTER, False), ('过程评价30%-50%，权重总和100%', WD_ALIGN_PARAGRAPH.LEFT, False), ('\u2705 过程40%+期末50%+思政10%', WD_ALIGN_PARAGRAPH.LEFT, False)],
        [('8. 教改摘要四维', WD_ALIGN_PARAGRAPH.CENTER, False), ('立德树人/思想/方法/手段均填写', WD_ALIGN_PARAGRAPH.LEFT, False), ('\u2705 四维均已按要求填写', WD_ALIGN_PARAGRAPH.LEFT, False)],
    ],
    [2.5, 5.5, 6.6]
)
doc.add_h2('五、补充说明')
doc.add_body('由于用户表示暂无可用的2026版培养方案.docx文件，本大纲按照智能科学与技术专业通用培养标准及同类高校该课程的通行设置编制。课程编码\uff08ZB2304012\uff09为参考性编码，实际使用时建议以学校教务处核定的编码为准。')

# Save
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, '数据结构与算法 课程教学大纲.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')

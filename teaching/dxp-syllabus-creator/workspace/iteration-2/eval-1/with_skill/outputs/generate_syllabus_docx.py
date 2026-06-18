# -*- coding: utf-8 -*-
"""
课程教学大纲 docx 生成脚本
用法：将本脚本复制到工作目录，在生成的 Python 脚本中 import 使用。
提供所有格式化功能，无需手动处理排版细节。
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


class SyllabusDocx:
    """教学大纲文档生成器，封装所有排版参数"""

    # 排版常量（按学校规范）
    FONT_NAME = '宋体'
    FONT_SIZE_BODY = Pt(12)       # 小四
    FONT_SIZE_H1 = Pt(16)         # 三号
    FONT_SIZE_H2 = Pt(14)         # 四号
    FONT_SIZE_TITLE = Pt(18)      # 小二
    LINE_SPACING = 1.25
    INDENT_FIRST_LINE = Cm(0.74)  # 首行缩进2字符
    HEADER_SHADE = 'D5E8F0'

    # 页面（A4）
    PAGE_W = Cm(21.0)
    PAGE_H = Cm(29.7)
    MARGIN_T = Cm(2.5)
    MARGIN_B = Cm(2.5)
    MARGIN_L = Cm(3.2)
    MARGIN_R = Cm(3.2)

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_default_style()

    def _setup_page(self):
        s = self.doc.sections[0]
        s.page_width = self.PAGE_W
        s.page_height = self.PAGE_H
        s.top_margin = self.MARGIN_T
        s.bottom_margin = self.MARGIN_B
        s.left_margin = self.MARGIN_L
        s.right_margin = self.MARGIN_R

    def _setup_default_style(self):
        style = self.doc.styles['Normal']
        style.font.name = self.FONT_NAME
        style.font.size = self.FONT_SIZE_BODY
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
        pf = style.paragraph_format
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _run(self, text, bold=False, size=None, name=None):
        run = self._p.add_run(text)
        run.bold = bold
        run.font.size = size or self.FONT_SIZE_BODY
        run.font.name = name or self.FONT_NAME
        run.element.rPr.rFonts.set(qn('w:eastAsia'), name or self.FONT_NAME)
        return run

    def _set_shading(self, cell, color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _new_para(self, spacing_before=0, spacing_after=3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        self._p = self.doc.add_paragraph()
        self._p.paragraph_format.line_spacing = self.LINE_SPACING
        self._p.paragraph_format.space_before = Pt(spacing_before)
        self._p.paragraph_format.space_after = Pt(spacing_after)
        self._p.alignment = alignment
        return self._p

    def add_title(self, text):
        """文档标题（居中、小二加粗）"""
        self._new_para(spacing_after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._run(text, bold=True, size=self.FONT_SIZE_TITLE)

    def add_h1(self, text):
        """一级标题（三号加粗）"""
        self._new_para(spacing_before=12, spacing_after=6)
        self._run(text, bold=True, size=self.FONT_SIZE_H1)

    def add_h2(self, text):
        """二级标题（四号加粗）"""
        self._new_para(spacing_before=8, spacing_after=4)
        self._run(text, bold=True, size=self.FONT_SIZE_H2)

    def add_h3(self, text):
        """三级标题（小四加粗）"""
        self._new_para(spacing_before=6, spacing_after=3)
        self._run(text, bold=True)

    def add_body(self, text):
        """正文段落（首行缩进2字符）"""
        self._new_para()
        self._p.paragraph_format.first_line_indent = self.INDENT_FIRST_LINE
        self._run(text)

    def add_body_no_indent(self, text):
        """正文段落（无缩进）"""
        self._new_para()
        self._run(text)

    def add_body_mixed(self, parts):
        """混合格式正文，parts = [(text, bold), ...]"""
        self._new_para()
        self._p.paragraph_format.first_line_indent = self.INDENT_FIRST_LINE
        for text, bold in parts:
            self._run(text, bold=bold)

    def add_blank_line(self):
        """空行"""
        self._new_para()
        self._p.paragraph_format.space_before = Pt(0)
        self._p.paragraph_format.space_after = Pt(0)
        self._run('')

    def add_page_break(self):
        """分页符"""
        self.doc.add_page_break()

    def add_table(self, headers, rows, col_widths_cm):
        """
        添加格式化表格
        :param headers: [str, ...] 表头文本列表
        :param rows: [[(text, alignment, bold), ...], ...] 数据行
        :param col_widths_cm: [float, ...] 每列宽度（厘米）
        """
        ncols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=ncols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 列宽
        for ci, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

        # 表头
        for ci, h in enumerate(headers):
            cell = table.rows[0].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = self.FONT_SIZE_BODY
            run.font.name = self.FONT_NAME
            run.element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
            self._set_shading(cell, self.HEADER_SHADE)

        # 数据行
        for ri, row_data in enumerate(rows):
            for ci, (text, align, bold) in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = align
                run = p.add_run(text)
                run.bold = bold
                run.font.size = self.FONT_SIZE_BODY
                run.font.name = self.FONT_NAME
                run.element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME)

        self.add_blank_line()
        return table

    def save(self, path):
        """保存文档"""
        self.doc.save(path)

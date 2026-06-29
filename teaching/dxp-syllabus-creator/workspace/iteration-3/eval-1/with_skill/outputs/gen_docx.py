# -*- coding: utf-8 -*-
"""
生成《数据结构与算法》课程教学大纲 docx
基于 示例：机械原理 课程教学大纲.docx 模板进行内容替换
"""

import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

FONT_NAME = '宋体'
FONT_SIZE = Pt(12)
GRADE_LIST = ["优 90-100", "良 80-90", "中 70-80", "及格 60-70", "不及格 0-60"]

WORK_DIR = r"C:\Users\admin\.agents\skills\dxp-syllabus-creator\workspace\iteration-3\eval-1\with_skill\outputs"
TEMPLATE_PATH = os.path.join(WORK_DIR, "DS202601-数据结构与算法 课程大纲 模板.docx")
OUTPUT_PATH = os.path.join(WORK_DIR, "DS202601-数据结构与算法 课程大纲[26版].docx")


def fix_font(run):
    """Set font to Song Ti"""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=FONT_SIZE):
    """Clear and set cell text with formatting"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    if p.runs:
        run = p.runs[0]
        run.text = text
        run.bold = bold
        run.font.size = size
        run.font.name = FONT_NAME
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    else:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = size
        run.font.name = FONT_NAME
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def global_replace(doc, old_text, new_text):
    """Replace text in all paragraphs and tables"""
    for para in doc.paragraphs:
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)


def set_column_width(table, col_index, width_cm):
    """Set width for a column in cm"""
    for row in table.rows:
        row.cells[col_index].width = Cm(width_cm)


def main():
    doc = docx.Document(TEMPLATE_PATH)

    # ============================================================
    # STEP 1: Global replaces (机械原理 -> 数据结构与算法, etc.)
    # ============================================================
    global_replace(doc, "机械原理", "数据结构与算法")
    global_replace(doc, "机械工程", "智能科学与技术")
    global_replace(doc, "机械设计制造及其自动化", "智能科学与技术")

    # ============================================================
    # STEP 2: Header table (Table 0) - 9 rows x 11 cols
    # Follow the skill rules: protect label cells, only fill value cells
    # ============================================================
    t0 = doc.tables[0]

    # R0: 课程编码 | value(C1) | 课程名称 | value(C2-C3) | ... | 课程性质 | value(C5-C6) | 课程负责人 | value(C7-C8) | 制表 | value(C10)
    set_cell_text(t0.rows[0].cells[1], "DS202601")  # 课程编码 value
    set_cell_text(t0.rows[0].cells[2], "数据结构与算法")  # 课程名称 value (merged C2-C3)
    # C4 is label "专业或大类/课程性质" - protect, C5-C6 merged for value
    set_cell_text(t0.rows[0].cells[5], "学科基础必修课")  # 课程性质 value
    # C7-C8 merged for 课程负责人 value
    set_cell_text(t0.rows[0].cells[7], "课程组")  # 课程负责人
    # C10 is 制表 value
    set_cell_text(t0.rows[0].cells[9], "课程组")  # 制表

    # R1: 课程名称全称 - C1 merged full width
    set_cell_text(t0.rows[1].cells[1], "数据结构与算法", align=WD_ALIGN_PARAGRAPH.LEFT)

    # R2: 英文名称 - C1 merged full width
    set_cell_text(t0.rows[2].cells[1], "Data Structures and Algorithms", align=WD_ALIGN_PARAGRAPH.LEFT)

    # R3: 学分 label(C0) | value(C1) | ... | 总学时 label(C3) | label(C4) | 讲课 label(C5) | 实验 label(C6) | label(C7) | 上机 label(C8) | label(C9) | 实践 label(C10)
    set_cell_text(t0.rows[3].cells[1], "3.5")  # 学分 value

    # R4: 执行学期 | value(C1) | ... | 总学时 value(C3) | ... | 讲课 value(C5) | 实验 value(C6) | ... | 上机 value(C8) | ... | 实践 value(C10)
    set_cell_text(t0.rows[4].cells[1], "第3学期", align=WD_ALIGN_PARAGRAPH.LEFT)  # 执行学期
    set_cell_text(t0.rows[4].cells[3], "64")    # 总学时
    set_cell_text(t0.rows[4].cells[5], "48")    # 讲课
    set_cell_text(t0.rows[4].cells[6], "0")     # 实验
    set_cell_text(t0.rows[4].cells[8], "16")    # 上机
    set_cell_text(t0.rows[4].cells[10], "0")    # 实践

    # R5: 考核方式 - full width value in C1
    set_cell_text(t0.rows[5].cells[1], "过程考核（平时作业、课堂测验、上机实验）50% + 期末闭卷考试50%", align=WD_ALIGN_PARAGRAPH.LEFT)

    # R6: 授课对象 - full width value in C1
    set_cell_text(t0.rows[6].cells[1], "智能科学与技术专业 大二本科生（第3学期）", align=WD_ALIGN_PARAGRAPH.LEFT)

    # R7: 先修课程 - full width value in C1
    set_cell_text(t0.rows[7].cells[1], "高级语言程序设计（C语言）、离散数学", align=WD_ALIGN_PARAGRAPH.LEFT)

    # R8: 后续课程 - full width value in C1
    set_cell_text(t0.rows[8].cells[1], "操作系统、数据库原理、算法设计与分析、机器学习、数据挖掘、计算机视觉、自然语言处理", align=WD_ALIGN_PARAGRAPH.LEFT)

    # ============================================================
    # STEP 3: Replace paragraph text content
    # ============================================================
    # Since we already did global replace for "机械原理", now replace specific paragraphs
    # by finding and replacing key phrases

    # Find and replace section titles if needed
    for para in doc.paragraphs:
        text = para.text.strip()

        # Replace title
        if text.startswith("机械原理") or "机械原理" in text and para.runs and para.runs[0].bold:
            for run in para.runs:
                if "机械原理" in run.text:
                    run.text = run.text.replace("机械原理", "数据结构与算法")
                if "宋体：四号黑体、加粗" in run.text:
                    run.text = run.text.replace("宋体：四号黑体、加粗", "")

    # ============================================================
    # STEP 4: Table 1 - 课程目标与毕业要求对应关系 (5 rows x 3 cols)
    # Header: 毕业要求 | 指标点 | 课程目标
    # ============================================================
    t1 = doc.tables[1]
    # Clear existing content rows (keep header row 0)
    for ri in range(1, len(t1.rows)):
        for ci in range(len(t1.columns)):
            for p in t1.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    # Row 1: 毕业要求1
    set_cell_text(t1.rows[1].cells[0], "1. 工程知识", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(t1.rows[1].cells[1], "1.1 掌握数学、自然科学、工程基础和智能科学专业知识，能够将其用于解决智能系统中的复杂工程问题", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(t1.rows[1].cells[2], "课程目标1（H）", align=WD_ALIGN_PARAGRAPH.LEFT)

    # Row 2: 毕业要求3
    set_cell_text(t1.rows[2].cells[0], "3. 设计/开发解决方案", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(t1.rows[2].cells[1], "3.1 能够针对智能科学与技术领域的复杂工程问题设计解决方案，开发满足特定需求的软硬件系统或模块", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(t1.rows[2].cells[2], "课程目标2（H）", align=WD_ALIGN_PARAGRAPH.LEFT)

    # Row 3: 毕业要求5
    set_cell_text(t1.rows[3].cells[0], "5. 使用现代工具", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(t1.rows[3].cells[1], "5.2 能够开发、选择与使用恰当的技术、资源和信息技术工具，对智能科学领域的复杂工程问题进行预测与模拟", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(t1.rows[3].cells[2], "课程目标3（M）", align=WD_ALIGN_PARAGRAPH.LEFT)

    # Row 4: clear if exists (template has 5 rows)
    if len(t1.rows) > 4:
        for ci in range(len(t1.columns)):
            for p in t1.rows[4].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    # ============================================================
    # STEP 5: Table 2 - 教学内容与课程目标对应关系 (6 rows x 3 cols)
    # Header: 课程目标 | 课程教学内容 | 学时
    # ============================================================
    t2 = doc.tables[2]
    chapter_data = [
        ("课程目标1", "第1章 绪论（数据结构基本概念、算法与算法分析）", "4"),
        ("课程目标1", "第2章 线性表（线性表的顺序与链式存储、基本操作实现）", "8"),
        ("课程目标1", "第3章 栈与队列（栈和队列的存储结构与基本操作、应用）", "8"),
        ("课程目标1", "第4章 串、数组与广义表（串的模式匹配、矩阵压缩存储）", "6"),
    ]

    for ri in range(1, len(t2.rows)):
        for ci in range(len(t2.columns)):
            for p in t2.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, (obj, content, hours) in enumerate(chapter_data):
        if ri + 1 < len(t2.rows):
            set_cell_text(t2.rows[ri + 1].cells[0], obj, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t2.rows[ri + 1].cells[1], content, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t2.rows[ri + 1].cells[2], hours)

    # ============================================================
    # STEP 6: Table 3 - 平时作业与课程目标对应关系 (6 rows x 4 cols)
    # Header: 课程目标 | 教学内容 | 平时作业内容 | 所占比例
    # ============================================================
    t3 = doc.tables[3]
    homework_data = [
        ("课程目标1", "第2章 线性表", "作业1：线性表顺序存储与链式存储操作", "20%"),
        ("课程目标1", "第3章 栈与队列", "作业2：栈与队列基本操作与应用", "20%"),
        ("课程目标2", "第5章 树与二叉树", "作业3：二叉树遍历与哈夫曼编码", "30%"),
        ("课程目标2", "第6章 图", "作业4：图的存储与遍历算法", "15%"),
        ("课程目标3", "第7-8章 查找与排序", "作业5：查找与排序算法综合分析", "15%"),
    ]

    for ri in range(1, len(t3.rows)):
        for ci in range(len(t3.columns)):
            for p in t3.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, (obj, chap, content, ratio) in enumerate(homework_data):
        if ri + 1 < len(t3.rows):
            set_cell_text(t3.rows[ri + 1].cells[0], obj, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t3.rows[ri + 1].cells[1], chap, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t3.rows[ri + 1].cells[2], content, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t3.rows[ri + 1].cells[3], ratio)

    # ============================================================
    # STEP 7: Table 4 - 平时作业评分标准 (16 rows x 3 cols)
    # Col: 课程目标 | 评价等级 | 评分标准
    # ============================================================
    t4 = doc.tables[4]
    hw_grading = [
        ("课程目标1", "优 90-100", "作业按时提交，数据结构基本概念理解准确，算法手工模拟完全正确，解答规范完整。"),
        ("课程目标1", "良 80-90", "作业按时提交，基本概念理解正确，算法模拟基本正确，存在少量细节不足。"),
        ("课程目标1", "中 70-80", "作业基本按时提交，部分概念理解有偏差，算法模拟存在少量错误。"),
        ("课程目标1", "及格 60-70", "作业偶有迟交，概念理解不完整，算法模拟存在较多错误。"),
        ("课程目标1", "不及格 0-60", "作业缺交或大面积错误，基本概念混淆，无法正确模拟算法过程。"),
        ("课程目标2", "优 90-100", "算法设计思路清晰新颖，数据结构选择恰当，代码实现功能完整规范，有性能优化思考。"),
        ("课程目标2", "良 80-90", "算法设计合理，数据结构选择正确，代码功能正确，规范性良好。"),
        ("课程目标2", "中 70-80", "算法设计基本合理，数据结构选择尚可，代码存在少量功能缺陷。"),
        ("课程目标2", "及格 60-70", "算法设计思路不清晰，数据结构选择不够合理，代码实现有明显缺陷。"),
        ("课程目标2", "不及格 0-60", "算法设计思路混乱，无法正确选择数据结构，代码无法运行。"),
        ("课程目标3", "优 90-100", "具有明显的算法效率优化意识，编码风格严谨规范，分析报告深入有条理。"),
        ("课程目标3", "良 80-90", "有一定的效率优化意识，编码较为规范，分析报告内容合理。"),
        ("课程目标3", "中 70-80", "效率分析基本正确，编码基本规范，报告分析较为简略。"),
        ("课程目标3", "及格 60-70", "效率分析存在较多不足，编码规范性差，报告内容不完整。"),
        ("课程目标3", "不及格 0-60", "缺乏效率分析意识，编码混乱，未提交分析报告。"),
    ]

    for ri in range(1, len(t4.rows)):
        for ci in range(len(t4.columns)):
            for p in t4.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, (obj, grade, desc) in enumerate(hw_grading):
        if ri + 1 < len(t4.rows):
            set_cell_text(t4.rows[ri + 1].cells[0], obj, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t4.rows[ri + 1].cells[1], grade, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t4.rows[ri + 1].cells[2], desc, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Set column widths for grading tables
    set_column_width(t4, 0, 2.5)   # 课程目标
    set_column_width(t4, 1, 2.5)   # 评价等级
    set_column_width(t4, 2, 9.0)   # 评分标准

    # ============================================================
    # STEP 8: Table 5 - 课堂测验与课程目标对应关系 (4 rows x 4 cols)
    # Header: 课程目标 | 教学内容 | 测验内容描述 | 所占比例
    # ============================================================
    t5 = doc.tables[5]
    quiz_data = [
        ("课程目标1", "第1-4章", "测验1：数据结构基础概念与线性结构", "50%"),
        ("课程目标1", "第5-8章", "测验2：树、图、查找与排序核心概念", "50%"),
    ]

    for ri in range(1, len(t5.rows)):
        for ci in range(len(t5.columns)):
            for p in t5.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, (obj, chap, content, ratio) in enumerate(quiz_data):
        if ri + 1 < len(t5.rows):
            set_cell_text(t5.rows[ri + 1].cells[0], obj, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t5.rows[ri + 1].cells[1], chap, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t5.rows[ri + 1].cells[2], content, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t5.rows[ri + 1].cells[3], ratio)

    # ============================================================
    # STEP 9: Table 6 - 课堂测验评分标准 (16 rows x 3 cols)
    # ============================================================
    t6 = doc.tables[6]
    quiz_grading = [
        ("课程目标1", "优 90-100", "数据结构核心概念理解准确深刻，经典算法执行过程模拟完全正确，答题规范。"),
        ("课程目标1", "良 80-90", "核心概念理解正确，算法模拟基本正确，存在少量细节偏差。"),
        ("课程目标1", "中 70-80", "基本概念理解存在部分偏差，算法模拟有少量步骤错误。"),
        ("课程目标1", "及格 60-70", "概念掌握不全面，算法模拟步骤缺失或混淆，答题不够规范。"),
        ("课程目标1", "不及格 0-60", "基本概念严重混淆，无法正确模拟算法执行过程。"),
    ]

    for ri in range(1, len(t6.rows)):
        for ci in range(len(t6.columns)):
            for p in t6.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, (obj, grade, desc) in enumerate(quiz_grading):
        if ri + 1 < len(t6.rows):
            set_cell_text(t6.rows[ri + 1].cells[0], obj, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t6.rows[ri + 1].cells[1], grade, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t6.rows[ri + 1].cells[2], desc, align=WD_ALIGN_PARAGRAPH.LEFT)

    set_column_width(t6, 0, 2.5)
    set_column_width(t6, 1, 2.5)
    set_column_width(t6, 2, 9.0)

    # ============================================================
    # STEP 10: Table 7 - 实验(上机)与课程目标对应关系 (4 rows x 4 cols)
    # Header: 教学目标 | 教学内容 | 实验内容 | 所占比例
    # ============================================================
    t7 = doc.tables[7]
    lab_data = [
        ("课程目标2", "第2-4章 线性结构", "上机1：线性表、栈与队列、串的基本操作实现", "37.5%"),
        ("课程目标2", "第5-6章 树与图", "上机2：二叉树遍历与哈夫曼编码、图的存储与遍历", "37.5%"),
        ("课程目标3", "第7-8章 查找与排序", "上机3：二叉排序树与哈希表、快速排序与堆排序实现", "25%"),
    ]

    for ri in range(1, len(t7.rows)):
        for ci in range(len(t7.columns)):
            for p in t7.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, (obj, chap, content, ratio) in enumerate(lab_data):
        if ri + 1 < len(t7.rows):
            set_cell_text(t7.rows[ri + 1].cells[0], obj, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t7.rows[ri + 1].cells[1], chap, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t7.rows[ri + 1].cells[2], content, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t7.rows[ri + 1].cells[3], ratio)

    # ============================================================
    # STEP 11: Table 8 - 实验评分标准 (6 rows x 3 cols)
    # ============================================================
    t8 = doc.tables[8]
    lab_grading = [
        ("课程目标2", "优 90-100", "程序功能完整正确，算法实现高效规范，代码注释清晰，实验报告分析深入，有创造性拓展。"),
        ("课程目标2", "良 80-90", "程序功能正确，算法实现合理，代码较规范，实验报告分析合理。"),
        ("课程目标2", "中 70-80", "程序基本功能正确，算法实现有少量缺陷，代码可读性一般，报告分析基本完整。"),
        ("课程目标2", "及格 60-70", "程序存在部分功能缺陷，算法实现不够完整，代码规范性差，报告分析简略。"),
        ("课程目标2", "不及格 0-60", "程序无法运行或严重偏离实验要求，算法实现混乱，实验报告缺失或质量极差。"),
    ]

    for ri in range(1, len(t8.rows)):
        for ci in range(len(t8.columns)):
            for p in t8.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, (obj, grade, desc) in enumerate(lab_grading):
        if ri + 1 < len(t8.rows):
            set_cell_text(t8.rows[ri + 1].cells[0], obj, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t8.rows[ri + 1].cells[1], grade, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t8.rows[ri + 1].cells[2], desc, align=WD_ALIGN_PARAGRAPH.LEFT)

    set_column_width(t8, 0, 2.5)
    set_column_width(t8, 1, 2.5)
    set_column_width(t8, 2, 9.0)

    # ============================================================
    # STEP 12: Table 9 - 期末考试与课程目标对应关系 (5 rows x 4 cols)
    # Header: 课程目标 | 教学内容 | 期末考试内容 | 所占比例
    # ============================================================
    t9 = doc.tables[9]
    exam_data = [
        ("课程目标1", "第1-8章", "选择题与填空题：数据结构基本概念、算法性质判断", "30%"),
        ("课程目标1", "第2-6章", "简答题：数据结构操作过程描述、算法原理阐述", "20%"),
        ("课程目标2", "第2-8章", "综合应用题与编程题：算法设计、数据结构选择与编程实现", "40%"),
        ("课程目标3", "第7-8章", "综合应用题：不同算法方案比较与效率分析", "10%"),
    ]

    for ri in range(1, len(t9.rows)):
        for ci in range(len(t9.columns)):
            for p in t9.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, (obj, chap, content, ratio) in enumerate(exam_data):
        if ri + 1 < len(t9.rows):
            set_cell_text(t9.rows[ri + 1].cells[0], obj, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t9.rows[ri + 1].cells[1], chap, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t9.rows[ri + 1].cells[2], content, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t9.rows[ri + 1].cells[3], ratio)

    # ============================================================
    # STEP 13: Table 10 - 期末考试评分标准 (16 rows x 3 cols)
    # ============================================================
    t10 = doc.tables[10]
    exam_grading = [
        ("课程目标1", "优 90-100", "数据结构概念理解深刻准确，经典算法原理阐述清晰完整，答题逻辑严谨。"),
        ("课程目标1", "良 80-90", "概念理解正确，算法原理阐述基本完整，存在少量不准确之处。"),
        ("课程目标1", "中 70-80", "核心概念基本理解，算法原理阐述有部分遗漏或偏差。"),
        ("课程目标1", "及格 60-70", "概念掌握不全面，算法原理阐述不够完整，存在较多错误。"),
        ("课程目标1", "不及格 0-60", "基本概念严重混淆，无法正确阐述算法原理。"),
        ("课程目标2", "优 90-100", "数据结构选择恰当，算法设计创新高效，编程实现完整正确，代码规范优雅。"),
        ("课程目标2", "良 80-90", "数据结构选择合理，算法设计正确，编程实现基本正确，代码较规范。"),
        ("课程目标2", "中 70-80", "数据结构选择基本合理，算法设计有小缺陷，编程实现存在部分问题。"),
        ("课程目标2", "及格 60-70", "数据结构选择不够优化，算法设计有较明显缺陷，编程实现不完整。"),
        ("课程目标2", "不及格 0-60", "无法正确选择数据结构，算法设计思路混乱，编程实现无法表达算法意图。"),
        ("课程目标3", "优 90-100", "算法效率分析准确深入，能对不同方案进行比较论证，思维缜密。"),
        ("课程目标3", "良 80-90", "算法效率分析基本准确，能够进行一定的方案比较。"),
        ("课程目标3", "中 70-80", "效率分析存在少量疏漏，方案比较不够全面。"),
        ("课程目标3", "及格 60-70", "效率分析有较多错误，缺乏方案比较意识。"),
        ("课程目标3", "不及格 0-60", "无法进行基本的算法效率分析，缺乏比较思维能力。"),
    ]

    for ri in range(1, len(t10.rows)):
        for ci in range(len(t10.columns)):
            for p in t10.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, (obj, grade, desc) in enumerate(exam_grading):
        if ri + 1 < len(t10.rows):
            set_cell_text(t10.rows[ri + 1].cells[0], obj, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t10.rows[ri + 1].cells[1], grade, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(t10.rows[ri + 1].cells[2], desc, align=WD_ALIGN_PARAGRAPH.LEFT)

    set_column_width(t10, 0, 2.5)
    set_column_width(t10, 1, 2.5)
    set_column_width(t10, 2, 9.0)

    # ============================================================
    # STEP 14: Table 11 - 课程目标达成度评价成绩构成 (7 rows x 7 cols)
    # ============================================================
    t11 = doc.tables[11]
    # Header row stays; we update data rows
    achievement_data = [
        ("1. 工程知识", "课程目标1", "平时作业 15%", "课堂测验 10%", "上机实验 5%", "期末考试 30%", "60%"),
        ("3. 设计/开发解决方案", "课程目标2", "-", "-", "上机实验 15%", "期末考试 40%", "55%"),
        ("5. 使用现代工具", "课程目标3", "-", "-", "上机实验 5%", "期末考试 10%", "15%"),
        ("合计", "", "15%", "10%", "25%", "80%", "100%"),
    ]

    for ri in range(1, 7):  # Rows 1-6
        for ci in range(len(t11.columns)):
            for p in t11.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.text = ''

    for ri, row_data in enumerate(achievement_data):
        if ri + 1 < len(t11.rows):
            for ci, val in enumerate(row_data):
                set_cell_text(t11.rows[ri + 1].cells[ci], val, align=WD_ALIGN_PARAGRAPH.LEFT,
                              bold=(ri == len(achievement_data) - 1))

    # ============================================================
    # STEP 15: Set column widths for remaining tables
    # ============================================================
    for ci, w in enumerate([3.5, 5.5, 5.0]):
        set_column_width(t1, ci, w)

    for ci, w in enumerate([3.0, 8.0, 2.0]):
        set_column_width(t2, ci, w)

    for ci, w in enumerate([2.5, 3.5, 6.0, 2.0]):
        set_column_width(t3, ci, w)
        set_column_width(t9, ci, w)

    for ci, w in enumerate([3.0, 5.0, 4.0, 2.0]):
        set_column_width(t5, ci, w)
        set_column_width(t7, ci, w)

    for ci, w in enumerate([3.0, 3.0, 3.5, 3.5, 3.5, 3.5, 2.5]):
        set_column_width(t11, ci, w)

    # ============================================================
    # SAVE
    # ============================================================
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == '__main__':
    main()

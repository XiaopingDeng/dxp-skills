# -*- coding: utf-8 -*-
"""
计算机网络 课程简介 docx 生成脚本
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT_NAME = '宋体'

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.2)
section.right_margin = Cm(3.2)

# Set default style
style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
style.paragraph_format.line_spacing = 1.25

def add_labeled_paragraph(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    run_label = p.add_run(label)
    run_label.bold = bold_label
    run_label.font.size = Pt(12)
    run_label.font.name = FONT_NAME
    run_label.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run_value = p.add_run(value)
    run_value.font.size = Pt(12)
    run_value.font.name = FONT_NAME
    run_value.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("计算机网络  课程简介")
run.bold = True
run.font.size = Pt(14)
run.font.name = FONT_NAME
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

# Blank line
doc.add_paragraph()

# Course Name
add_labeled_paragraph(doc, "课程名称：", "计算机网络")

# Course Code
add_labeled_paragraph(doc, "课程编码：", "CN202602")

# Course Introduction
intro_text = (
    "计算机网络是智能科学与技术专业的一门专业必修课，3学分，52学时（其中理论讲授44学时，实验8学时），在第5学期开设。"
    "本课程以计算机网络的体系结构为主线，系统讲授物理层、数据链路层、网络层、运输层和应用层的基本概念、基本原理和关键技术。"
    "课程内容涵盖OSI和TCP/IP参考模型、数据通信基础、信道复用技术、循环冗余检验、CSMA/CD协议、以太网交换技术、"
    "IP地址规划与子网划分、CIDR无分类编址、RIP和OSPF路由选择协议、TCP可靠传输与拥塞控制、HTTP/DNS/FTP等应用层协议、"
    "以及Wireshark协议分析和Packet Tracer网络仿真等实验技能。"
    "通过本课程的学习，学生能够掌握计算机网络的设计思想和协议机制，具备网络拓扑分析、协议报文解析和网络应用设计的基本能力，"
    "为后续学习物联网技术、云计算与大数据处理、网络与信息安全等课程以及从事网络应用开发和智能系统集成工作奠定坚实的网络知识基础。"
    "本课程支撑毕业要求2（问题分析，H）、要求3（设计/开发解决方案，M）和要求5（使用现代工具，H）的达成。"
)
add_labeled_paragraph(doc, "课程简介：", intro_text)

# Assessment Method
add_labeled_paragraph(doc, "考核形式：", "考试（过程考核占50%，含平时作业20%、课堂测验15%、实验15%；期末考试占50%）")

# Recommended Textbooks and References
ref_text = (
    "建议教材：谢希仁.《计算机网络（第8版）》. 电子工业出版社, 2021."
    "参考书目：[1] James F. Kurose, Keith W. Ross. Computer Networking: A Top-Down Approach (8th Edition). Pearson, 2020. "
    "[2] Andrew S. Tanenbaum, David J. Wetherall. Computer Networks (6th Edition). Pearson, 2020. "
    "[3] 吴功宜.《计算机网络（第4版）》. 清华大学出版社, 2019."
)
add_labeled_paragraph(doc, "建议教材与参考书目：", ref_text)

output_path = r"C:\Users\admin\.agents\skills\dxp-syllabus-creator\workspace\iteration-3\eval-2\with_skill\outputs\CN202602-计算机网络 课程简介[26版].docx"
doc.save(output_path)
print(f"Saved: {output_path}")

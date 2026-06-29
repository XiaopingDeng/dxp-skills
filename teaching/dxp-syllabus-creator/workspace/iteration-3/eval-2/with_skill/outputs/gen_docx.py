# -*- coding: utf-8 -*-
"""
计算机网络 课程教学大纲 docx 生成脚本
基于"示例：机械原理 课程教学大纲.docx"模板替换内容生成
"""
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

FONT_NAME = '宋体'
FONT_SIZE = Pt(12)
GRADES = ["优 90-100", "良 80-90", "中 70-80", "及格 60-70", "不及格 0-60"]

doc = docx.Document(r"C:\Users\admin\.agents\skills\dxp-syllabus-creator\workspace\iteration-3\eval-2\with_skill\outputs\CN202602-计算机网络 课程大纲 模板.docx")

def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12)):
    """安全设置单元格文本"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def clear_para(p):
    """清空段落"""
    for r in p.runs:
        r.text = ''
    for t_elem in p._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        t_elem.text = ''

def fill_para(para, text, bold=False, size=Pt(12), font_name=FONT_NAME):
    """填充段落文本"""
    clear_para(para)
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run("")
    run.text = text
    run.bold = bold
    run.font.size = size
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# ============================================================
# 全局文本替换（将模板中机械原理相关文本替换为计算机网络内容）
# ============================================================
for p in doc.paragraphs:
    for run in p.runs:
        if "机械原理" in run.text:
            run.text = run.text.replace("机械原理", "计算机网络")
        if "机械工程" in run.text:
            run.text = run.text.replace("机械工程", "智能科学与技术")

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if "机械原理" in run.text:
                        run.text = run.text.replace("机械原理", "计算机网络")
                    if "机械工程" in run.text:
                        run.text = run.text.replace("机械工程", "智能科学与技术")

# ============================================================
# Table 0: 页眉信息表 (9行 x 11列)
# ============================================================
t0 = doc.tables[0]

# R0: 课程编码 | CN202602 | 课程名称 | 计算机网络 | ... | ... | ... | 课程负责人 | ... | ... |
set_cell(t0.rows[0].cells[1], "CN202602")   # 课程编码值
# C2 is merged with C3 for 课程名称
set_cell(t0.rows[0].cells[2], "计算机网络")
# C5: 课程性质 (merged C4-C6), label in C4
set_cell(t0.rows[0].cells[5], "专业必修课")
# C7: 课程负责人 (merged C7-C8)
set_cell(t0.rows[0].cells[7], "")
# C9: label is "制表", C10=制表人
set_cell(t0.rows[0].cells[9], "")

# R1: 课程名称 (全宽)
set_cell(t0.rows[1].cells[1], "计算机网络", align=WD_ALIGN_PARAGRAPH.LEFT)

# R2: 英文名称 (全宽)
set_cell(t0.rows[2].cells[1], "Computer Networks", align=WD_ALIGN_PARAGRAPH.LEFT)

# R3: 学分
set_cell(t0.rows[3].cells[1], "3")

# R4: 执行学期 | 第5学期 | | 总学时 | 52 | 讲课 | 44 | 实验 | 8 | 上机 | 0 | 实践 | 0
set_cell(t0.rows[4].cells[1], "第5学期", align=WD_ALIGN_PARAGRAPH.LEFT)
set_cell(t0.rows[4].cells[3], "52")   # 总学时
set_cell(t0.rows[4].cells[5], "44")   # 讲课
set_cell(t0.rows[4].cells[6], "8")    # 实验
set_cell(t0.rows[4].cells[8], "0")    # 上机
set_cell(t0.rows[4].cells[10], "0")   # 实践

# R5: 考核方式
set_cell(t0.rows[5].cells[1], "考试（过程考核50%+期末考试50%）", align=WD_ALIGN_PARAGRAPH.LEFT)

# R6: 授课对象
set_cell(t0.rows[6].cells[1], "智能科学与技术专业本科生", align=WD_ALIGN_PARAGRAPH.LEFT)

# R7: 先修课程
set_cell(t0.rows[7].cells[1], "C语言程序设计基础，数据结构与算法，计算机组成原理", align=WD_ALIGN_PARAGRAPH.LEFT)

# R8: 后续课程
set_cell(t0.rows[8].cells[1], "物联网技术，云计算与大数据处理，网络与信息安全，分布式系统", align=WD_ALIGN_PARAGRAPH.LEFT)

# ============================================================
# 正文段落替换
# ============================================================
paras = doc.paragraphs

# P0: 标题 - 计算机网络 课程教学大纲
fill_para(paras[0], "计算机网络  课程教学大纲", bold=True, size=Pt(16))

# P1-P3: Section header and intro
fill_para(paras[1], "一、课程的性质、任务和目标", bold=True, size=Pt(14))
fill_para(paras[2], "", bold=False, size=Pt(12))
fill_para(paras[3], "1. 课程性质")

# P4: 课程性质正文
fill_para(paras[4], "本课程是智能科学与技术专业的一门专业必修课，是该专业学生构建计算机网络知识体系的核心课程。课程以计算机网络的体系结构为主线，系统讲授物理层、数据链路层、网络层、运输层和应用层的基本概念、基本原理和关键技术，使学生掌握计算机网络的设计思想和协议机制，为后续学习分布式系统、云计算、物联网等课程以及从事网络应用开发和智能系统集成奠定坚实基础。")

# P5: 课程任务 header
fill_para(paras[5], "2. 课程任务")

# P6: 课程任务正文
fill_para(paras[6], "本课程的主要任务是通过系统讲授计算机网络的体系结构、各层协议的工作原理及网络互联技术，使学生理解数据在网络中的传输过程，掌握TCP/IP协议族核心协议的功能与交互机制，具备网络拓扑分析、协议报文解析和网络应用设计的基本能力。结合课程思政，将家国情怀、科技自立自强意识融入教学全过程，引导学生树立正确的网络安全观和科技伦理观。")

# P7: 课程目标 header
fill_para(paras[7], "3. 课程目标")

# P8: 课程目标1
fill_para(paras[8], "课程目标1：掌握计算机网络的体系结构（OSI和TCP/IP参考模型），理解物理层、数据链路层、网络层、运输层和应用层的基本概念、核心协议及工作原理。能够运用网络分层思想分析复杂网络工程问题，识别和判断通信过程中的关键环节与性能瓶颈，支撑毕业要求2（问题分析，H）。")

# P9: 课程目标2
fill_para(paras[9], "课程目标2：掌握IP地址规划与子网划分、路由选择算法、拥塞控制机制、可靠传输原理等网络核心技术与方法，能够针对具体的网络通信需求设计合理的网络拓扑与协议配置方案，初步具备网络系统设计与开发的能力，支撑毕业要求3（设计/开发解决方案，M）。")

# P10: 课程目标3
fill_para(paras[10], "课程目标3：熟练使用Wireshark等网络协议分析工具进行数据包捕获与协议分析，掌握Cisco Packet Tracer或GNS3等网络仿真平台进行网络拓扑设计与协议验证，能够利用现代工具和仿真环境进行网络性能测试与故障诊断，支撑毕业要求5（使用现代工具，H）。")

# P11: 课程目标与毕业要求对应的说明
fill_para(paras[11], "课程目标与毕业要求的对应关系如表1所示。")

# P12: 表1标签
fill_para(paras[12], "表1 课程目标与毕业要求的对应关系")

# ============================================================
# Table 1: 课程目标与毕业要求对应关系 (5行 x 3列)
# ============================================================
t1 = doc.tables[1]
# Row 0: Headers are already fine (毕业要求 | 指标点 | 课程目标)
# Clear old data rows and fill with new data
t1_data = [
    ["2. 问题分析\n（H）", "2.1 能够运用数学、自然科学和工程科学的基本原理，识别和判断复杂智能系统工程问题的关键环节与核心要素", "课程目标1"],
    ["3. 设计/开发解决方案\n（M）", "3.1 能够针对复杂智能系统工程问题，设计满足特定需求的解决方案或系统模块", "课程目标2"],
    ["5. 使用现代工具\n（H）", "5.1 能够选择与使用合适的网络协议分析工具、仿真平台等现代工具，对复杂智能系统工程问题进行模拟与预测", "课程目标3"],
]

for ri in range(1, min(len(t1.rows), len(t1_data) + 1)):
    row = t1.rows[ri]
    set_cell(row.cells[0], t1_data[ri-1][0], bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[1], t1_data[ri-1][1], bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(row.cells[2], t1_data[ri-1][2], bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# 第二节 课程教学内容
# ============================================================
fill_para(paras[13], "二、课程教学内容的基本要求、重点和难点及学时分配")
fill_para(paras[14], "第1章 计算机网络概述（4学时）")
fill_para(paras[15], "基本要求")
fill_para(paras[16], "（1）了解计算机网络在信息时代中的作用与发展历程，理解互联网的标准化工作。\n（2）掌握计算机网络的组成与分类，理解电路交换、报文交换和分组交换的基本原理与区别。\n（3）掌握计算机网络的性能指标（速率、带宽、吞吐量、时延、时延带宽积、往返时间RTT、利用率）。\n（4）掌握计算机网络的体系结构——OSI七层模型与TCP/IP四层模型，理解协议与服务的概念。")
fill_para(paras[17], "重点和难点")
fill_para(paras[18], "重点是分组交换原理和TCP/IP体系结构；难点是网络性能指标的综合计算与协议分层思想的深刻理解。")
fill_para(paras[19], "教学方法")
fill_para(paras[20], "理论讲授结合网络拓扑案例演示，辅以互联网发展历史中关键事件的技术解读。")

# P21-P22 空行
fill_para(paras[21], "")
fill_para(paras[22], "")

# P23: 表2标签
fill_para(paras[23], "表2 课程教学内容与课程目标的对应关系及学时分配")

# P24: 后续需要处理剩余章节，但模板段落数有限，我们在已有结构基础上继续
# 模板中有大量段落用于机械原理的分章内容，我们使用段落27-28之间的空行
fill_para(paras[24], "第2章 物理层（6学时）")
fill_para(paras[25], "基本要求")
fill_para(paras[26], "（1）理解物理层的基本概念和主要任务，掌握数据通信的基础知识（信道的极限容量、奈奎斯特定理和香农公式）。\n（2）掌握物理层下面的传输媒体（双绞线、同轴电缆、光纤、无线信道）的特性与应用场景。\n（3）理解信道复用技术（频分复用FDM、时分复用TDM、波分复用WDM、码分复用CDM）的基本原理。\n（4）了解宽带接入技术（ADSL、光纤同轴混合网HFC、FTTx）。")

# Check total paragraphs
total_paras = len(paras)
print(f"Total paragraphs: {total_paras}")

# Map remaining paragraphs for sections 2-6 content
# P27 through P64 are from the template, we need to carefully remap

# Due to template structure limitations, we'll insert critical content via paragraphs
# and handle the rest with direct overwrites where possible

# ============================================================
# Table 2: 课程教学内容与课程目标对应 (template has 6 rows)
# ============================================================
t2 = doc.tables[2]
t2_data = [
    ["课程目标1", "第1章 计算机网络概述", "4"],
    ["课程目标1", "第2章 物理层", "6"],
    ["课程目标1", "第3章 数据链路层（含实验1：2学时）", "8"],
    ["课程目标1/2/3", "第4章 网络层（含实验2：3学时）", "12"],
    ["课程目标1/2/3", "第5章 运输层（含实验3：3学时）", "8"],
]
# Ensure table headers are right
t2_header = t2.rows[0]
set_cell(t2_header.cells[0], "课程目标", bold=True)
set_cell(t2_header.cells[1], "教学内容", bold=True)
set_cell(t2_header.cells[2], "学时", bold=True)

for ri in range(1, min(len(t2.rows), len(t2_data) + 1)):
    row = t2.rows[ri]
    set_cell(row.cells[0], t2_data[ri-1][0], bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[1], t2_data[ri-1][1], bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(row.cells[2], t2_data[ri-1][2], bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# Process remaining paragraphs (P27 onward aren't all mappable from template)
# Since template has specific structure for机械原理, we replace content
# ============================================================

# The template paragraphs 27 onwards contain remaining sections.
# Let's overwrite section headers and key content paragraphs

# Find and replace section headers
section_mapping = {
    "三、建议教材与学习资源": "29",  # approximate index
    "四、考核形式与成绩构成": "39",
    "五、有关说明": "65",
    "六、课程建设与改革摘要": "68",
}

# Overwrite section 3 header
for i, p in enumerate(paras):
    txt = p.text.strip()
    if "建议教材" in txt and "学习资源" in txt:
        fill_para(p, "三、建议教材与学习资源", bold=True, size=Pt(14))
    elif "考核形式" in txt and "成绩构成" in txt:
        fill_para(p, "四、考核形式与成绩构成", bold=True, size=Pt(14))
    elif txt.startswith("五") and ("说明" in txt or "有关" in txt):
        fill_para(p, "五、有关说明", bold=True, size=Pt(14))
    elif "课程建设" in txt and "改革" in txt:
        fill_para(p, "六、课程建设与改革摘要", bold=True, size=Pt(14))

# ============================================================
# Replace 建议教材 section content
# The template has specific paragraphs for 建议教材/参考教材/学习资源
# ============================================================
for i, p in enumerate(paras):
    t = p.text.strip()
    if "建议教材" == t or t.startswith("建议教材"):
        fill_para(p, "建议教材")
        # Next para should be the textbook
        if i+1 < len(paras):
            fill_para(paras[i+1], "[1] 谢希仁.《计算机网络（第8版）》. 电子工业出版社, 2021.（\"十二五\"普通高等教育本科国家级规划教材）")
    elif t == "学习资源" or t.startswith("学习资源"):
        fill_para(p, "参考教材")
        if i+1 < len(paras):
            fill_para(paras[i+1], "[1] James F. Kurose, Keith W. Ross. Computer Networking: A Top-Down Approach (8th Edition). Pearson, 2020.\n[2] Andrew S. Tanenbaum, David J. Wetherall. Computer Networks (6th Edition). Pearson, 2020.\n[3] 吴功宜.《计算机网络（第4版）》. 清华大学出版社, 2019.\n[4] 王卫东, 傅宇凡.《计算机网络习题解析与同步辅导》. 电子工业出版社, 2022.")
    elif "线上" in t and ("资源" in t or "课程" in t):
        fill_para(p, "线上学习资源")
        if i+1 < len(paras):
            fill_para(paras[i+1], "[1] 中国大学MOOC——《计算机网络》国家精品课程（华南理工大学），https://www.icourse163.org/course/SCUT-1206672813\n[2] MIT OpenCourseWare — 6.829 Computer Networks, https://ocw.mit.edu/courses/6-829-computer-networks-fall-2002/\n[3] Stanford Online — CS 144: Introduction to Computer Networking, https://cs144.github.io/\n[4] 国家高等教育智慧教育平台——计算机网络课程资源库，https://www.chinaooc.com.cn/")
    elif "数智" in t or "数字" in t or "AI" in t:
        fill_para(p, "数智化教学资源")
        if i+1 < len(paras):
            fill_para(paras[i+1], "[1] 网络协议分析工具：Wireshark（开源软件，用于各层协议数据包捕获与分析）\n[2] 网络仿真平台：Cisco Packet Tracer / GNS3（用于网络拓扑设计与路由协议仿真验证）\n[3] AI辅助学习工具：ChatGPT / 智谱清言（用于协议原理问答互动与疑难知识点解析）\n[4] 在线编程练习平台：LeetCode（网络相关算法题）/ PTA（程序设计类实验辅助教学平台）")

# ============================================================
# Replace 考核形式 section content
# ============================================================
for i, p in enumerate(paras):
    t = p.text.strip()
    if "成绩构成" in t or (t.startswith("1") and "成绩" in t):
        fill_para(p, "1. 成绩构成")
        if i+1 < len(paras) and ("课程成绩" in paras[i+1].text or "%" in paras[i+1].text):
            fill_para(paras[i+1], "课程总成绩由过程考核成绩（占50%）和期末考试成绩（占50%）两部分组成。过程考核包括平时作业（20%）、课堂测验（15%）和实验（15%）三项。")
    elif "过程考核" in t or (t.startswith("2") and "过程" in t):
        fill_para(p, "2. 过程考核")
    elif "期末" in t or (t.startswith("3") and "期末" in t):
        fill_para(p, "3. 期末考试")
        # Find next para to put exam description
        if i+1 < len(paras):
            fill_para(paras[i+1], "期末考试采用闭卷笔试形式，满分100分，占总成绩50%。试题涵盖全部教学内容，题型包括选择题（约20%）、填空题（约10%）、简答题（约30%）、计算与分析题（约40%），注重考查学生对网络协议原理的理解和运用能力，以及对网络问题的分析与解决能力。")

# ============================================================
# Replace 五、有关说明 content
# ============================================================
for i, p in enumerate(paras):
    t = p.text.strip()
    if t.startswith("五") and ("说明" in t or "有关" in t):
        # Next paragraph should be the detailed content
        if i+1 < len(paras):
            fill_para(paras[i+1], "1. 先修衔接：本课程以《C语言程序设计基础》中的基本编程能力和《数据结构与算法》中的逻辑思维训练为必备基础，以《计算机组成原理》中的计算机硬件体系结构与数据表示知识为前置支撑。学生须具备C语言编程调试能力和计算机硬件基础知识方可顺利修读本课程。\n\n2. 后续承接：本课程为后续《物联网技术》《云计算与大数据处理》《网络与信息安全》《分布式系统》《软件定义网络》《智能嵌入式系统设计》等课程提供计算机网络体系结构、TCP/IP协议族、网络通信机制等核心知识支撑。\n\n3. 开课学期建议：本课程在第5学期开设。此时学生已在前4学期完成程序设计、数据结构、计算机组成原理等前置课程的学习，具备较好的编程基础和计算机系统认知能力。\n\n4. 修读建议：本课程理论性强、概念密集，学生须在课堂学习之外每周投入不少于4小时的课外学习时间（含教材精读、习题练习和实验准备）。建议学生充分利用Wireshark等工具对日常网络通信进行实际观察以加深理解，利用Packet Tracer进行自主组网实验以提升实践能力。\n\n5. 教学调整机制：如遇教学内容更新需要调整章节学时分配，由任课教师在开课前提出调整方案，经课程组讨论通过后报系教学主任审批。\n\n6. 免修条件：获得HCIA及以上或CCNA及以上网络相关认证证书，或在省级及以上网络技术竞赛中获得二等奖及以上奖励的学生，可申请免修。\n\n7. 学术诚信要求：实验报告和作业必须独立完成，如有抄袭行为，当次作业或实验成绩记为0分，并按学校相关学术诚信管理规定处理。")
        break

# ============================================================
# Replace 六、课程建设与改革摘要 content
# ============================================================
section6_added = False
for i, p in enumerate(paras):
    t = p.text.strip()
    if t.startswith("六") and ("课程建设" in t or "改革" in t):
        section6_added = True
        # We need to add subsections after this
        # Due to template structure, we'll insert content via nearby paragraphs
        continue

# Find paragraphs after section 6 header to fill with 4 dimensions
# Template may have limited paragraphs; let's check what's available
# P68-P71 are around section 6 in template

# Let's target specific paragraphs near the end
section6_paras = []
for i in range(total_paras):
    t = paras[i].text.strip()
    if "立德树人" in t or "教学思想" in t or "教学方法" in t or "教学手段" in t:
        section6_paras.append(i)

section6_content = {
    "立德树人": "本课程牢牢把握立德树人根本任务，深度挖掘计算机网络课程中的思政元素与科学精神，将思想政治教育有机融入专业知识教学中。在讲授计算机网络发展史时，通过介绍中国在5G通信标准制定中的核心贡献（如华为Polar码入选5G控制信道编码方案），培养学生的民族自豪感和科技报国情怀；在讲授网络安全内容时，结合\"没有网络安全就没有国家安全\"的重要论断，引导学生树立正确的网络安全观和科技伦理意识；在网络协议标准制定的教学内容中，强调中国从标准跟随者到标准引领者的转变历程，激发学生自主创新、攻坚克难的奋斗精神。通过实验环节的团队协作要求，培养学生的合作意识、责任担当和职业道德素养。",
    "教学思想": "本课程坚持以学生为中心、以产出为导向、以质量为核心的教学理念，系统设计课程目标、教学内容与评价方式的反向映射关系。以毕业要求指标点（问题分析H、设计/开发方案M、使用现代工具H）为顶层导向，细化课程目标与教学内容的对应矩阵，确保每项教学活动都有明确的产出指向。学生中心理念体现于：根据学生前置知识基础和个体差异进行分层教学，基础薄弱学生配备额外辅导和补充材料，能力突出学生提供进阶阅读（如RFC文档研读）和挑战任务。面对AI时代人才培养新需求，本课程在原有知识传授基础上，强化AI赋能网络的前沿视角，引导学生理解人工智能技术对网络工程范式的影响。通过\"知识传授+能力培养+价值引领\"三位一体的教学设计，实现学生核心素养与关键能力的全面提升。",
    "教学方法": "本课程构建\"理论+实践+创新\"三维教学模式，综合运用多种教学方法以适应不同教学内容的特点。（1）理论讲授+案例教学：用于概念原理密集的教学内容（如TCP/IP体系结构、各层协议工作原理），以实际网络通信场景为案例切入点，将抽象协议转化为可感知的通信过程。（2）问题驱动+研讨互动：用于重点难点内容（如子网划分、路由算法、拥塞控制），以典型工程问题为驱动，组织课堂分组讨论，使学生在争议和辨析中深化理解。（3）项目驱动实验教学：用于实验环节（Wireshark抓包分析、Packet Tracer网络设计、TCP行为观测），学生以完成具体实验任务为目标，在实际操作中验证理论并发现问题。（4）翻转课堂：用于应用层部分章节（如HTTP/HTTPS协议、DNS体系），学生课前通过视频与文献自主学习基础知识，课堂时间用于深度研讨和案例分析。",
    "教学手段": "本课程深入推进教育数字化转型，依托智慧教学平台与AI工具构建全流程智能化教学生态。（1）智慧教学平台：使用超星学习通作为课程主平台，承载课件发布、作业提交与批改、在线测验、讨论互动等功能；利用雨课堂进行课堂实时互动（弹幕提问、随堂测试、签到考勤），提升课堂参与度。（2）AI教学工具：引入Wireshark协议分析工具和Cisco Packet Tracer/GNS3网络仿真环境作为核心实验平台；推荐学生使用大语言模型（如智谱清言、Kimi）作为个性化学习助手，用于协议原理的交互式问答和疑难知识点解析。（3）数字化教学资源：整合中国大学MOOC平台国家精品课程视频、CS 144（Stanford）公开课视频、RFC标准文档等优质数字化教学资源，构建\"教材+MOOC+公开课+技术文档\"四位一体的资源体系。（4）\"课前-课中-课后\"闭环管理：课前通过学习通发布预习任务和自测题；课中利用雨课堂实时互动收集学情数据，动态调整教学节奏；课后通过学习通布置分层作业，利用平台自动批改和统计分析功能实现精准反馈。教师定期分析教学过程数据，精准识别学习困难学生并给予个性化指导，以数智赋能推动教学质量的持续提升。",
}

for i, p in enumerate(paras):
    t = p.text.strip()
    for dim_name, dim_content in section6_content.items():
        if dim_name in t:
            # This is a sub-header, fill it and the next paragraph
            fill_para(p, f"{dim_name}")
            if i + 1 < len(paras) and len(paras[i+1].text.strip()) > 10:
                fill_para(paras[i+1], dim_content)

# ============================================================
# Table 3: 平时作业与课程目标对应关系 (6行 x 4列)
# ============================================================
t3 = doc.tables[3]
t3_data = [
    ["课程目标1", "第1-2章", "作业1：体系结构与性能指标计算", "20"],
    ["课程目标1", "第2章", "作业2：信道容量与复用技术", "20"],
    ["课程目标1", "第3章", "作业3：CRC校验与CSMA/CD分析", "20"],
    ["课程目标2", "第4章", "作业4：IP地址规划与子网划分", "20"],
    ["课程目标2", "第5章", "作业5：TCP拥塞控制与可靠传输", "20"],
]
for ri in range(1, min(len(t3.rows), len(t3_data) + 1)):
    row = t3.rows[ri]
    set_cell(row.cells[0], t3_data[ri-1][0], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[1], t3_data[ri-1][1], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[2], t3_data[ri-1][2], align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(row.cells[3], t3_data[ri-1][3], align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# Table 4: 平时作业评分标准 (16行 x 3列)
# ============================================================
t4 = doc.tables[4]
# 5 grades each for 3 course objectives = 15 data rows + header
# Template has rows for 课程目标1(作业1-2), 课程目标2(作业3-4), 课程目标3(作业5)
# We adapt: 课程目标1->作业1-2, 课程目标2->作业3-4, 课程目标2->作业5 (since we only have 2 targets for homework)
hw_rubric = {
    "课程目标1\n（作业1、2）": [
        "完全掌握网络体系结构的分层思想，正确绘制OSI和TCP/IP参考模型，准确计算网络性能指标（速率、时延、RTT等），正确计算信道容量并选择适当的复用技术。",
        "较好掌握网络体系结构，能基本正确地绘制参考模型，能较准确计算网络性能指标和信道容量，复用技术选择基本恰当。",
        "基本理解网络体系结构，能画出大致分层结构，能进行简单的性能指标计算，但对综合性计算和复用技术原理理解有偏差。",
        "对网络体系结构理解不够完整，分层模型绘制有较多错误，性能指标计算存在较多问题，对复用技术理解较为模糊。",
        "未能理解网络体系结构的分层思想，不能正确绘制参考模型，不能完成性能指标计算和复用技术分析。",
    ],
    "课程目标2\n（作业3、4、5）": [
        "完全掌握CRC校验计算方法，正确分析CSMA/CD协议的争用期和碰撞检测过程；准确完成IP地址的子网划分和CIDR聚合计算，正确分析TCP拥塞控制各阶段窗口变化。",
        "较好掌握CRC计算和CSMA/CD分析，能较准确完成子网划分和CIDR计算，能基本正确分析TCP拥塞控制过程，存在少量计算错误。",
        "基本掌握CRC计算方法和CSMA/CD概念，能进行基本子网划分但对CIDR聚合理解有偏差，能描述TCP拥塞控制算法但对具体窗口计算有一定困难。",
        "对CRC计算和CSMA/CD原理理解不完整，子网划分存在较多错误，TCP拥塞控制分析仅能定性描述，不能完成定量计算。",
        "未能掌握CRC计算方法，不能分析CSMA/CD协议工作过程，不能完成IP子网划分和TCP拥塞控制分析。",
    ],
}

# We have 3 objectives but only 2 groups in homework
# Course goal 1 has rows 1-5 (5 grades), course goal 2 has rows 6-10 and 11-15
hw_obj_order = ["课程目标1\n（作业1、2）", "课程目标2\n（作业3、4）", "课程目标2\n（作业5）"]
obj_idx = 0
for ri in range(1, len(t4.rows)):
    row = t4.rows[ri]
    if (ri - 1) % 5 == 0 and (ri - 1) // 5 < len(hw_obj_order):
        obj_idx = (ri - 1) // 5
    grade_idx = (ri - 1) % 5

    if obj_idx < len(hw_obj_order):
        set_cell(row.cells[0], hw_obj_order[obj_idx], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row.cells[1], GRADES[grade_idx], align=WD_ALIGN_PARAGRAPH.CENTER)

        # Get rubric text based on which objective group
        if obj_idx == 0:
            rubric_key = "课程目标1\n（作业1、2）"
        else:
            rubric_key = "课程目标2\n（作业3、4、5）"
        set_cell(row.cells[2], hw_rubric[rubric_key][grade_idx], align=WD_ALIGN_PARAGRAPH.LEFT)

# ============================================================
# Table 5: 课堂测验与课程目标对应关系 (4行 x 4列)
# ============================================================
t5 = doc.tables[5]
t5_data = [
    ["课程目标1", "第1-2章", "测验1：概述+物理层", "20"],
    ["课程目标1/2", "第3-4章", "测验2：数据链路层+网络层", "60"],
    ["课程目标2", "第5-6章", "测验3：运输层+应用层", "20"],
]
for ri in range(1, min(len(t5.rows), len(t5_data) + 1)):
    row = t5.rows[ri]
    set_cell(row.cells[0], t5_data[ri-1][0], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[1], t5_data[ri-1][1], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[2], t5_data[ri-1][2], align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(row.cells[3], t5_data[ri-1][3], align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# Table 6: 课堂测验评分标准 (16行 x 3列)
# ============================================================
t6 = doc.tables[6]
quiz_rubric = {
    "课程目标1\n（测验1）": [
        "完全掌握计算机网络概念与分类，正确理解分组交换原理，准确计算各类网络性能指标，掌握OSI和TCP/IP模型各层功能，准确运用香农公式进行信道容量分析。",
        "较好掌握网络基本概念，正确理解分组交换，能较准确计算网络性能指标，能正确描述OSI和TCP/IP模型，基本掌握香农公式应用。",
        "基本理解网络概念和分组交换原理，能进行简单的性能指标计算，能描述OSI和TCP/IP模型的主要层次，对香农公式应用存在部分困难。",
        "对网络基本概念理解不够完整，对分组交换与电路交换的区分模糊，性能指标计算存在较多错误，对体系结构的分层理解不清晰。",
        "未能掌握计算机网络基本概念，不理解分组交换原理，不能计算网络性能指标，不能描述网络体系结构。",
    ],
    "课程目标2\n（测验2、3）": [
        "完全掌握CRC和CSMA/CD的工作原理与计算，正确完成任意复杂度的子网划分和CIDR聚合，准确掌握RIP和OSPF路由算法及路由表生成过程，正确分析TCP连接管理和拥塞控制全过程。",
        "较好掌握数据链路层和网络层核心协议原理，能较准确完成子网划分和路由表分析，基本掌握TCP拥塞控制算法，存在少量计算错误。",
        "基本理解数据链路层和网络层主要协议功能，能进行基本子网划分，能描述路由算法原理，能定性描述TCP拥塞控制过程。",
        "对数据链路层和网络层协议理解不完整，子网划分存在较多错误，对路由算法理解模糊，TCP拥塞控制分析能力较弱。",
        "未能掌握数据链路层和网络层核心协议，不能完成子网划分，不能理解路由算法，不能分析TCP拥塞控制。",
    ],
}

quiz_obj_order = ["课程目标1\n（测验1）", "课程目标2\n（测验2、3）", "课程目标2\n（测验2、3）"]
q_obj_idx = 0
for ri in range(1, len(t6.rows)):
    row = t6.rows[ri]
    if (ri - 1) % 5 == 0 and (ri - 1) // 5 < len(quiz_obj_order):
        q_obj_idx = (ri - 1) // 5
    grade_idx = (ri - 1) % 5

    if q_obj_idx < len(quiz_obj_order):
        set_cell(row.cells[0], quiz_obj_order[q_obj_idx], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row.cells[1], GRADES[grade_idx], align=WD_ALIGN_PARAGRAPH.CENTER)
        if q_obj_idx == 0:
            rubric_key = "课程目标1\n（测验1）"
        else:
            rubric_key = "课程目标2\n（测验2、3）"
        set_cell(row.cells[2], quiz_rubric[rubric_key][grade_idx], align=WD_ALIGN_PARAGRAPH.LEFT)

# ============================================================
# Table 7: 实验与课程目标对应关系 (4行 x 4列)
# ============================================================
t7 = doc.tables[7]
t7_data = [
    ["课程目标3", "实验1", "实验1：网络协议分析与以太网帧格式观测（2学时）", "33.3"],
    ["课程目标2/3", "实验2", "实验2：IP地址规划与路由协议配置仿真（3学时）", "33.3"],
    ["课程目标2/3", "实验3", "实验3：TCP协议分析与拥塞控制观测（3学时）", "33.4"],
]
for ri in range(1, min(len(t7.rows), len(t7_data) + 1)):
    row = t7.rows[ri]
    set_cell(row.cells[0], t7_data[ri-1][0], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[1], t7_data[ri-1][1], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[2], t7_data[ri-1][2], align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(row.cells[3], t7_data[ri-1][3], align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# Table 8: 实验评分标准 (6行 x 3列)
# ============================================================
t8 = doc.tables[8]
lab_rubric = [
    "能够独立完成实验操作，熟练使用Wireshark/Packet Tracer等工具，准确捕获和分析协议数据包，正确完成网络拓扑配置与路由协议设置，实验报告内容完整、数据详实、分析深入。",
    "能够较独立地完成实验操作，较熟练地使用工具，较准确地完成协议分析和网络配置，实验报告内容较完整、分析较合理。",
    "能够在教师指导下完成实验操作，基本掌握工具使用方法，能完成基本的协议分析和网络配置，实验报告内容基本完整但分析深度不足。",
    "仅能在教师帮助下完成部分实验操作，工具使用不够熟练，协议分析和网络配置存在较多问题，实验报告内容不够完整。",
    "无法独立完成实验操作，不能正确使用工具，不能完成协议分析和网络配置，未提交实验报告或报告质量极差。",
]
for ri in range(1, len(t8.rows)):
    row = t8.rows[ri]
    set_cell(row.cells[0], f"课程目标2/3\n（实验1、2、3）", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[1], GRADES[ri-1], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[2], lab_rubric[ri-1], align=WD_ALIGN_PARAGRAPH.LEFT)

# ============================================================
# Table 9: 期末考试与课程目标对应关系 (5行 x 4列)
# ============================================================
t9 = doc.tables[9]
t9_data = [
    ["课程目标1", "第1-2章", "计算机网络概述、物理层基础知识", "20"],
    ["课程目标1", "第3章", "数据链路层协议与机制", "16"],
    ["课程目标2", "第4章", "网络层IP规划与路由协议", "44"],
    ["课程目标2", "第5-6章", "运输层与应用层协议", "20"],
]
for ri in range(1, min(len(t9.rows), len(t9_data) + 1)):
    row = t9.rows[ri]
    set_cell(row.cells[0], t9_data[ri-1][0], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[1], t9_data[ri-1][1], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[2], t9_data[ri-1][2], align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(row.cells[3], t9_data[ri-1][3], align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# Table 10: 期末考试评分标准 (16行 x 3列)
# ============================================================
t10 = doc.tables[10]
exam_rubric = {
    "课程目标1": [
        "完全掌握计算机网络体系结构、物理层信道容量计算、数据链路层协议机制（CSMA/CD、CRC），正确回答所有基础概念题和分析题，答案准确完整。",
        "较好掌握体系结构和下层协议原理，能较准确回答基础概念题和分析题，存在少量概念表述不精确。",
        "基本掌握体系结构和下层协议主要概念，能回答大部分基础题，但对部分协议细节和分析题理解有偏差。",
        "对体系结构和下层协议理解不够完整，仅能回答简单概念题，分析题和计算题困难较大。",
        "未能掌握体系结构和下层协议基础知识，不能正确回答基础概念题。",
    ],
    "课程目标2": [
        "完全掌握IP地址规划、子网划分、CIDR聚合、RIP/OSPF路由算法，正确完成所有计算题和分析题，TCP拥塞控制和可靠传输分析准确无误。",
        "较好掌握网络层和运输层核心内容，能较准确完成计算题和分析题，存在少量计算偏差或分析不够深入。",
        "基本掌握IP规划和路由算法，能完成基本计算题，但对复杂子网划分和拥塞控制定量分析有一定困难。",
        "对IP规划和路由算法理解不完整，仅能完成简单计算题，复杂分析题和综合题困难较大。",
        "未能掌握IP规划、路由算法和运输层核心内容，不能完成计算题和分析题。",
    ],
}

exam_obj_order = ["课程目标1", "课程目标2", "课程目标2", "课程目标2"]
e_obj_idx = 0
for ri in range(1, len(t10.rows)):
    row = t10.rows[ri]
    if (ri - 1) % 5 == 0 and (ri - 1) // 5 < len(exam_obj_order):
        e_obj_idx = (ri - 1) // 5
    grade_idx = (ri - 1) % 5

    if e_obj_idx < len(exam_obj_order):
        set_cell(row.cells[0], exam_obj_order[e_obj_idx], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row.cells[1], GRADES[grade_idx], align=WD_ALIGN_PARAGRAPH.CENTER)
        if e_obj_idx == 0:
            rk = "课程目标1"
        else:
            rk = "课程目标2"
        set_cell(row.cells[2], exam_rubric[rk][grade_idx], align=WD_ALIGN_PARAGRAPH.LEFT)

# ============================================================
# Table 11: 课程目标达成度评价成绩构成 (7行 x 7列)
# ============================================================
t11 = doc.tables[11]
# Headers
t11_r0 = t11.rows[0]
set_cell(t11_r0.cells[0], "毕业要求指标点", bold=True)
set_cell(t11_r0.cells[1], "课程目标", bold=True)
set_cell(t11_r0.cells[6], "总比例", bold=True)

t11_r1 = t11.rows[1]
set_cell(t11_r1.cells[0], "毕业要求指标点", bold=True)
set_cell(t11_r1.cells[1], "课程目标", bold=True)
set_cell(t11_r1.cells[2], "平时作业", bold=True)
set_cell(t11_r1.cells[3], "课堂测验", bold=True)
set_cell(t11_r1.cells[4], "实验", bold=True)
set_cell(t11_r1.cells[5], "期末考试", bold=True)
set_cell(t11_r1.cells[6], "总比例", bold=True)

t11_data = [
    ["2.1（H）", "课程目标1", "8", "6", "", "18", "32"],
    ["3.1（M）", "课程目标2", "12", "9", "7", "22", "50"],
    ["5.1（H）", "课程目标3", "", "", "8", "10", "18"],
]
for ri in range(2, min(len(t11.rows), 2 + len(t11_data))):
    row = t11.rows[ri]
    d = t11_data[ri - 2]
    for ci, val in enumerate(d):
        set_cell(row.cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER)

# Last row: 合计
last_row = t11.rows[5]
set_cell(last_row.cells[0], "合计", bold=True)
set_cell(last_row.cells[1], "合计", bold=True)
set_cell(last_row.cells[2], "20", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell(last_row.cells[3], "15", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell(last_row.cells[4], "15", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell(last_row.cells[5], "50", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell(last_row.cells[6], "100", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# 调整表格列宽
# ============================================================
# Table 4 (评分标准) - col0: 2.5cm, col1: 2cm, col2: 10cm
for table_idx in [4, 6, 8, 10]:
    if table_idx < len(doc.tables):
        t = doc.tables[table_idx]
        for row in t.rows:
            row.cells[0].width = Cm(2.8)
            row.cells[1].width = Cm(2.2)
            row.cells[2].width = Cm(10.0)

# Table 11 (达成度) - adjust all columns
t = doc.tables[11]
widths = [Cm(2.2), Cm(2.0), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8)]
for row in t.rows:
    for ci, w in enumerate(widths):
        row.cells[ci].width = w

# ============================================================
# 页脚
# ============================================================
footer_targets = {"编 写 人：": "", "审 核 人：": "（教研室主任或课程组负责人）", "批 准 人：": "（学院院长）", "编写日期：": "2026 年  月  日"}
for i, p in enumerate(paras):
    t = p.text.strip()
    if "编 写" in t and "人" in t:
        fill_para(p, "编 写 人：")
    elif "审 核" in t and "人" in t:
        fill_para(p, "审 核 人：（教研室主任或课程组负责人）")
    elif "批 准" in t and "人" in t:
        fill_para(p, "批 准 人：（学院院长）")
    elif "编写日期" in t or "日期" in t:
        fill_para(p, "编写日期：2026 年  月  日")

# ============================================================
# 保存
# ============================================================
output_path = r"C:\Users\admin\.agents\skills\dxp-syllabus-creator\workspace\iteration-3\eval-2\with_skill\outputs\CN202602-计算机网络 课程大纲[26版].docx"
doc.save(output_path)
print(f"Saved: {output_path}")
print("Done!")

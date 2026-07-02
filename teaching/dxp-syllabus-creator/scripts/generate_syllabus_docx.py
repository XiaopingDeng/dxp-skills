# -*- coding: utf-8 -*-
"""
课程教学大纲 docx 生成脚本 — 基于模板的版本

两种用法：
  1. SyllabusFromTemplate(template_path) — 打开已有模板 docx，填充内容
  2. SyllabusDocx() — 从零创建文档（备用，不推荐）

模板结构说明：
  - 理论课程模板: 页眉表 9行x11列（含学分/总学时/讲课/实验/上机/实践明细）
  - 课程设计/实习/毕设模板: 页眉表 3行x4列（简洁型）
"""

import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================
# 常量
# ============================================================
FONT_NAME = '宋体'           # 中文字体（eastAsia）
WESTERN_FONT = 'Times New Roman'  # 西文字体（ascii）
HEITI_FONT = '黑体'          # 标题中文字体
FONT_SIZE = Pt(12)
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _make_run(run, text, bold=False, size=None, font_name=None, east_asia=None):
    """统一设置 run 属性。
    
    默认西文字体 Times New Roman、中文字体宋体，与参考模板一致。
    节标题等需用黑体时传 east_asia='黑体'。
    """
    run.text = text
    run.bold = bold
    run.font.size = size or FONT_SIZE
    run.font.name = font_name or WESTERN_FONT
    ea = east_asia or FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), ea)


def _set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格文本"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    _make_run(p.add_run(text), text, bold=bold)


def _shade(cell, color='D5E8F0'):
    """给单元格加底色"""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))


def _new_para_elem():
    """创建空白段落 XML 元素（西文 Times New Roman，中文宋体）"""
    return parse_xml(
        f'<w:p xmlns:w="{NS}"><w:r><w:rPr><w:rFonts w:ascii="{WESTERN_FONT}" w:eastAsia="{FONT_NAME}"/>'
        f'<w:sz w:val="24"/></w:rPr><w:t></w:t></w:r></w:p>')


# ============================================================
# SyllabusFromTemplate — 基于模板的文档生成器
# ============================================================
class SyllabusFromTemplate:
    """基于模板 docx 的大纲文档生成器。

    用法:
        doc = SyllabusFromTemplate("模板路径.docx")
        doc.fill_header({...})          # 填充页眉信息表
        doc.set_title("课程名称")       # 替换标题
        doc.set_section(0, "一、...")    # 替换第 i 节标题
        doc.insert_body(ref, "正文")     # 在 ref 后插入正文
        doc.insert_h3(ref, "子标题")     # 在 ref 后插入子标题
        doc.insert_table(ref, ...)       # 在 ref 后插入表格
        doc.save("输出路径.docx")
    """

    def __init__(self, template_path):
        self.doc = docx.Document(template_path)
        self._setup_style()
        # 页眉表 — 理论课是第0个表，其他也是第0个表
        self.header_table = self.doc.tables[0] if self.doc.tables else None
        self._detect_template_type()
        # 预定位各节标题段落（按模板固定索引）
        self._index_sections()

    def _setup_style(self):
        s = self.doc.styles['Normal']
        s.font.name = WESTERN_FONT
        s.font.size = FONT_SIZE
        s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        s.paragraph_format.line_spacing = 1.25

    def _detect_template_type(self):
        """识别模板类型"""
        if self.header_table and len(self.header_table.rows) == 9:
            self.template_type = 'theory'
        else:
            self.template_type = 'other'

    def _index_sections(self):
        """定位模板中各节标题和页脚的段落对象。

        理论模板段落索引（打开后）:
            P0=标题, P1=空, P2=一, P3=二, P4=三, P5=四, P6=五, P7=六,
            P8=空, P9=编写人, P10=审核人, P11=批准人, P12=日期, P13+=空

        设计/实习/毕设模板段落索引:
            P0=标题, P1=一, P2=二, P3=三, P4=四, P5=五, P6=六,
            P7=空, P8=编写人, P9=审核人, P10=批准人, P11=日期
        """
        paras = self.doc.paragraphs
        # 节标题: 非空且非标题的第1~6个段落
        section_paras = []
        footer_paras = []
        for p in paras:
            t = p.text.strip()
            if len(t) > 5 and ('�' not in t):  # non-garbled meaningful text
                if any(kw in t for kw in ['编 写', '审 核', '批 准', '编写日期']):
                    footer_paras.append(p)
                elif p is not paras[0]:
                    section_paras.append(p)

        # Fallback: use fixed indices based on template type
        if self.template_type == 'theory':
            self.sections = {
                0: paras[2],  # 一
                1: paras[3],  # 二
                2: paras[4],  # 三
                3: paras[5],  # 四
                4: paras[6],  # 五
                5: paras[7],  # 六
            }
            self.footer = {
                'author': paras[9] if len(paras) > 9 else None,
                'reviewer': paras[10] if len(paras) > 10 else None,
                'approver': paras[11] if len(paras) > 11 else None,
                'date': paras[12] if len(paras) > 12 else None,
            }
        else:
            self.sections = {
                0: paras[1],  # 一
                1: paras[2],  # 二
                2: paras[3],  # 三
                3: paras[4],  # 四
                4: paras[5],  # 五
                5: paras[6],  # 六
            }
            self.footer = {
                'author': paras[8] if len(paras) > 8 else None,
                'reviewer': paras[9] if len(paras) > 9 else None,
                'approver': paras[10] if len(paras) > 10 else None,
                'date': paras[11] if len(paras) > 11 else None,
            }

    # ── 页眉信息表填充 ──

    def fill_header_theory(self, code, name, en_name, credits, total_hours,
                           lecture_hours, lab_hours, practice_hours=0,
                           computer_hours=0, semester='', assessment='',
                           audience='', prerequisites='', followups='',
                           supervisor='', course_type=''):
        """填充理论课程页眉表（9行x11列）。

        表结构:
          R0: 课程编码 | value | 课程类别 | 课程类别 | 课程性质值(span) | | | 课程负责人 | 课程负责人 | 负责人值 | 负责人值
          R1: 课程名称 | value (全宽合并)
          R2: 英文名称 | value (全宽合并)
          R3: 学分 | value | | 总学时 | 总学时 | 讲课 | 实验 | 实验 | 上机 | 上机 | 实践
          R4: 执行学期 | value(span) | | value(总学时) | | value(讲) | value(实验) | | value(上机) | | value(实践)
          R5: 考核方式 | value (全宽合并)
          R6: 授课对象 | value (全宽合并)
          R7: 先修课程 | value (全宽合并)
          R8: 后续课程 | value (全宽合并)

        注意：
        - R0C7-C8 为"课程负责人"标签，禁止覆写。
        - R0C9-C10 为课程负责人值区，留给用户手工填写，禁止写入。
        """
        t = self.header_table
        if not t or len(t.rows) < 9:
            return

        set_c = _set_cell
        # R0 — 仅写入值单元格，不碰标签
        set_c(t.rows[0].cells[1], code)                    # 编码值
        set_c(t.rows[0].cells[4], course_type)              # 课程性质值（span=3）
        # R0C7-C10 课程负责人区域（含标签和值区）禁止写入 —— 跳过
        # R1 — 课程名称全称
        set_c(t.rows[1].cells[1], name, align=WD_ALIGN_PARAGRAPH.LEFT)
        # R2 — 英文名称
        set_c(t.rows[2].cells[1], en_name, align=WD_ALIGN_PARAGRAPH.LEFT)
        # R3 — 学分值
        set_c(t.rows[3].cells[1], str(credits))
        # R4 — 学时和执行学期
        set_c(t.rows[4].cells[1], semester, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_c(t.rows[4].cells[3], str(total_hours))
        set_c(t.rows[4].cells[5], str(lecture_hours))
        set_c(t.rows[4].cells[6], str(lab_hours))           # 实验（merged with C7）
        set_c(t.rows[4].cells[8], str(computer_hours))      # 上机（merged with C9）
        set_c(t.rows[4].cells[10], str(practice_hours))     # 实践
        # R5 — 考核方式
        set_c(t.rows[5].cells[1], assessment, align=WD_ALIGN_PARAGRAPH.LEFT)
        # R6 — 授课对象
        set_c(t.rows[6].cells[1], audience, align=WD_ALIGN_PARAGRAPH.LEFT)
        # R7 — 先修课程
        set_c(t.rows[7].cells[1], prerequisites, align=WD_ALIGN_PARAGRAPH.LEFT)
        # R8 — 后续课程
        set_c(t.rows[8].cells[1], followups, align=WD_ALIGN_PARAGRAPH.LEFT)

    def fill_header_other(self, code, name, en_name, supervisor='',
                          prerequisites='', period_semester=''):
        """填充设计/实习/毕设页眉表（3行x4列）。

        表结构:
          R0: 课程编码 | value | 课程负责人 | value
          R1: 课程名称（中文）| value | 课程名称（英文）| value
          R2: 先修课程 | value | 实践周期及执行学期 | value
        """
        t = self.header_table
        if not t or len(t.rows) < 3:
            return
        set_c = _set_cell
        set_c(t.rows[0].cells[1], code)
        set_c(t.rows[0].cells[3], supervisor)
        set_c(t.rows[1].cells[1], name, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_c(t.rows[1].cells[3], en_name, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_c(t.rows[2].cells[1], prerequisites, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_c(t.rows[2].cells[3], period_semester, align=WD_ALIGN_PARAGRAPH.LEFT)

    # ── 标题 ──

    def set_title(self, text):
        """替换文档标题（P0）。标题格式：课程名称 + 空格 + 课程教学大纲"""
        p = self.doc.paragraphs[0]
        self._clear_para(p)
        run = p.runs[0] if p.runs else p.add_run()
        _make_run(run, text, bold=True, east_asia=HEITI_FONT)

    # ── 节标题 ──

    def set_section(self, index, text):
        """替换第 index 节标题（0=第一节）。节标题使用黑体。"""
        p = self.sections.get(index)
        if p is None:
            return
        self._clear_para(p)
        run = p.runs[0] if p.runs else p.add_run()
        _make_run(run, text, bold=True, size=Pt(14), east_asia=HEITI_FONT)

    # ── 内容插入 ──

    def _clear_para(self, p):
        for r in p.runs:
            r.text = ''
        for t_elem in p._element.findall(f'{{{NS}}}t'):
            t_elem.text = ''

    def insert_body(self, ref_element, text, indent=True):
        """在指定元素后插入正文段落。返回新段落的 _element。"""
        p_elem = _new_para_elem()
        ref_element.addnext(p_elem)
        para = docx.text.paragraph.Paragraph(p_elem, self.doc)
        if not indent:
            para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        _make_run(run, text)
        return p_elem

    def insert_h3(self, ref_element, text):
        """在指定元素后插入子标题。返回新段落的 _element。"""
        p_elem = _new_para_elem()
        ref_element.addnext(p_elem)
        para = docx.text.paragraph.Paragraph(p_elem, self.doc)
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        _make_run(run, text, bold=True)
        return p_elem

    def insert_table(self, ref_element, headers, rows, col_widths_cm):
        """在指定元素后插入表格。返回表格的 _tbl 元素。

        Args:
            headers: [str, ...] 表头
            rows: [[(text, alignment), ...], ...] 数据行
            col_widths_cm: [float, ...] 列宽（厘米）
        """
        ncols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=ncols)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 表头
        for ci, h in enumerate(headers):
            cell = table.rows[0].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = FONT_SIZE
            run.font.name = WESTERN_FONT
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            _shade(cell)

        # 数据行
        for ri, row_data in enumerate(rows):
            for ci, (text, align) in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = align
                run = p.add_run(text)
                run.font.size = FONT_SIZE
                run.font.name = WESTERN_FONT
                run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                # 合计行加粗
                if ri == len(rows) - 1:
                    run.bold = True

        # 列宽
        for ci, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

        # 将表格移到 ref_element 后面
        ref_element.addnext(table._tbl)
        return table._tbl

    def insert_label(self, ref_element, text):
        """插入居中的表格标签行（如"表1 XXXX"）。返回新段落 _element。"""
        p_elem = _new_para_elem()
        ref_element.addnext(p_elem)
        para = docx.text.paragraph.Paragraph(p_elem, self.doc)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        _make_run(run, text)
        return p_elem

    # ── 页脚 ──

    def set_footer_author(self, text=''):
        """设置编写人"""
        p = self.footer.get('author')
        if p:
            self._clear_para(p)
            run = p.runs[0] if p.runs else p.add_run()
            _make_run(run, f'编 写 人：{text}')

    def set_footer_reviewer(self, text=''):
        """设置审核人"""
        p = self.footer.get('reviewer')
        if p:
            self._clear_para(p)
            run = p.runs[0] if p.runs else p.add_run()
            _make_run(run, f'审 核 人：{text}')

    def set_footer_approver(self, text=''):
        """设置批准人"""
        p = self.footer.get('approver')
        if p:
            self._clear_para(p)
            run = p.runs[0] if p.runs else p.add_run()
            _make_run(run, f'批 准 人：{text}')

    def set_footer_date(self, text='2026 年  月  日'):
        """设置编写日期"""
        p = self.footer.get('date')
        if p:
            self._clear_para(p)
            run = p.runs[0] if p.runs else p.add_run()
            _make_run(run, f'编写日期：{text}')

    # ── 保存 ──

    def save(self, path):
        self.doc.save(path)


# ============================================================
# 理论课程 docx 替换工具函数
# 这些函数用于"基于示例大纲替换内容"的策略
# 适用于理论课程（使用示例：机械原理模板）
# ============================================================

import copy
from docx.text.paragraph import Paragraph

# 已知的页眉表标签文字（遇到这些文字的单元格不覆写）
HEADER_LABELS = {
    '课程编码', '课程名称', '英文名称', '课程类别', '课程负责人',
    '学分', '总学时', '讲课', '实验', '上机', '实践',
    '考核方式', '执行学期', '学生对象', '授课对象', '先修课程', '后续课程',
}

# Footer 关键词（以这些开头的段落不替换）
FOOTER_KEYWORDS = ('编 写 人', '审 核 人', '批 准 人', '编写日期')

# 评分等级固定文本（4级制，与参考模板 XD26010006 一致）
GRADE_LEVELS = ["优\n90-100", "良\n80-89", "中/及格\n60-79", "差\n0-59"]


def add_run_songti(p, text, bold=False, size=Pt(10.5)):
    """向段落添加 run 并显式设置宋体字体（中文）+ Times New Roman（西文）。

    解决 p.clear() 后新 run 不继承模板字体的问题。
    参考模板使用双字体体系：西文 Times New Roman，中文宋体。
    所有正文段落替换都应使用此函数而非 p.add_run()。
    """
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.name = WESTERN_FONT
    run.font.size = size
    r_elem = run.element.rPr
    if r_elem is not None:
        rFonts = r_elem.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = r_elem.makeelement(qn('w:rFonts'), {})
            r_elem.append(rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
    return run


def set_cell_font(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10.5)):
    """安全设置单元格文本——先清空再写入，同时设置宋体字体。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    add_run_songti(p, text, size=size)


def insert_paragraph_before(ref_p, text, bold=False):
    """在 ref_p 段落之前插入新段落。

    用于在模板的"……"占位段落之前插入额外章节的
    基本要求/重点和难点/教学方法段落。

    原理：深拷贝参考段落的 XML 元素，清除内容，
    用 addprevious 插入到目标位置之前。
    """
    new_elem = copy.deepcopy(ref_p._element)
    # 清除所有 run 子元素
    for child in list(new_elem):
        if child.tag.endswith('}r'):
            new_elem.remove(child)
    ref_p._element.addprevious(new_elem)
    new_p = Paragraph(new_elem, ref_p._parent)
    add_run_songti(new_p, text, bold=bold)
    return new_p


def is_label_cell(cell):
    """检查单元格是否为标签单元格（不应覆写）。"""
    txt = cell.text.strip()
    return txt in HEADER_LABELS or any(label in txt for label in HEADER_LABELS)


def is_footer_para(p):
    """检查段落是否为 footer（不应替换）。"""
    txt = p.text.strip()
    return txt.startswith(FOOTER_KEYWORDS)


def safe_replace_course_name(doc, old_name, new_name):
    """安全全局替换课程名称。

    只替换完整课程名称（如"机械原理"→"新课程名"），
    绝不替换单字或部分名称（如"机械"→"计算机"），
    避免破坏出版社名称和模板结构术语。
    """
    for p in doc.paragraphs:
        for run in p.runs:
            if old_name in run.text:
                run.text = run.text.replace(old_name, new_name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if old_name in run.text:
                            run.text = run.text.replace(old_name, new_name)


def fill_header_table(table, data_map):
    """填充页眉表，自动保护标签单元格和课程负责人区域。

    data_map: {row_idx: {col_idx: value}} 格式的字典。
    遍历表格所有单元格，跳过标签单元格和"负责人"区域，
    仅填充 data_map 中指定的值单元格。
    """
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if is_label_cell(cell):
                continue
            # 跳过课程负责人区域（包含"负责人"文本的行/列）
            if '负责人' in cell.text:
                continue
            # 检查同行是否有"负责人"标签
            row_text = ' '.join(c.text.strip() for c in row.cells)
            if '负责人' in row_text and ci > 0:
                # 在负责人所在行，跳过值区
                continue
            # 填充值
            if ri in data_map and ci in data_map[ri]:
                set_cell_font(cell, data_map[ri][ci])


def fill_table_from_data(table, data, align_map=None):
    """用数据数组填充表格。

    data: 二维数组，data[0] 为表头行，data[1:] 为数据行。
    align_map: {col_idx: WD_ALIGN_PARAGRAPH} 指定每列对齐方式。
    """
    if align_map is None:
        align_map = {}
    for ri, row_data in enumerate(data):
        if ri >= len(table.rows):
            break
        for ci, val in enumerate(row_data):
            if ci >= len(table.rows[ri].cells):
                break
            align = align_map.get(ci, WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_font(table.rows[ri].cells[ci], val, align=align)


def clear_excess_rows(table, start_row):
    """清空表格中从 start_row 开始的多余数据行。

    当课程目标数/作业次数少于模板行数时，
    多余数据行的每个单元格必须清空，否则模板旧内容会残留。
    """
    for ri in range(start_row, len(table.rows)):
        for ci in range(len(table.rows[ri].cells)):
            table.rows[ri].cells[ci].text = ''


def replace_section_body(doc, heading_keyword, body_text):
    """找到包含 heading_keyword 的段落作为标题，
    替换其后的第一个段落为 body_text。

    使用文本特征匹配而非固定段落索引。
    跳过 footer 段落。
    """
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if is_footer_para(p):
            continue
        if heading_keyword in txt:
            # 替换标题后的第一个段落
            if i + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[i + 1]
                if not is_footer_para(next_p):
                    next_p.clear()
                    add_run_songti(next_p, body_text)
                    return True
    return False


def adjust_grading_table_widths(table, col0_cm=3.0, col1_cm=2.5, col2_cm=9.0):
    """调整评分标准表的列宽。"""
    for row in table.rows:
        if len(row.cells) >= 3:
            row.cells[0].width = Cm(col0_cm)
            row.cells[1].width = Cm(col1_cm)
            row.cells[2].width = Cm(col2_cm)


def adjust_data_table_widths(table, widths_cm):
    """调整数据对应表的列宽。

    widths_cm: [宽度列表]，如 [2.5, 2.5, 7.0, 2.0]
    """
    for row in table.rows:
        for ci, w in enumerate(widths_cm):
            if ci < len(row.cells):
                row.cells[ci].width = Cm(w)

